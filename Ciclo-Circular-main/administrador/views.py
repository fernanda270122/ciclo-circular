# ==============================================================
# IMPORTS COMPLETOS (Reemplaza el inicio de administrador/views.py)
# ==============================================================
import os
import json
import base64
import textwrap
from django.contrib.auth.tokens import default_token_generator
from django.utils.http import urlsafe_base64_encode
from django.utils.encoding import force_bytes
from django.urls import reverse
from django.utils.crypto import get_random_string
from django.conf import settings
from django.core.mail import EmailMessage, get_connection
from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect
import threading
from collections import Counter
from django.db.models import Q
from django.contrib import messages
from django.shortcuts import render, redirect
from django.contrib.auth.hashers import make_password
from django.core.mail import send_mail
from openpyxl import load_workbook
import random
import string
import random
import string
import cloudinary.uploader
import collections
import traceback # Importante para ver errores detallados
from io import BytesIO
from datetime import datetime


# --- Terceros ---
import numpy as np
import docx
import fitz # PyMuPDF
from wordcloud import WordCloud, STOPWORDS
from openpyxl import Workbook, load_workbook
from Levenshtein import distance, editops, apply_edit, jaro
from openai import OpenAI
from dotenv import load_dotenv
from xhtml2pdf import pisa # <--- Necesario para el PDF

try:
    from xhtml2pdf import pisa
except ImportError:
    pisa = None
    
# --- Django Core ---
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse, Http404
from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.hashers import make_password
from django.contrib.auth.models import User
from django.contrib.staticfiles import finders # <--- Necesario para imágenes en PDF
from django.template.loader import get_template
from django.views.generic import TemplateView
from django.core.mail import send_mail, EmailMessage, get_connection
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Q, Count # <--- Tu import solicitado

# --- Django Rest Framework (si lo usas) ---
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework import status


# --- Modelos Propios ---
from user.models import Usuario, TrabajoEmpresa, Preferencia
from app.models import (
    Universidad, Facultad, Departamento, Carrera,
    Etapa, RegistroActividad, Entrada, Salida, Oportunidades, 
    Idea, CVUsuario, Oferta, Necesidad,
    Evento, Invitacion, PreguntaEvento, 
    
)
from .models import LogTelegram

# Cargar variables de entorno
load_dotenv()

# Configuración IA (si aplica)
import google.generativeai as genai
try:
    genai.configure(api_key=settings.OPENAI_API_KEY)
except:
    pass

# ==============================================================
# FIN DE IMPORTS
# ==============================================================


# Create your views here.

def homeAdmin(request):
    # 1. Datos existentes
    registros = RegistroActividad.objects.filter(usuario=request.user)
    Universidads = Universidad.objects.all()

    # 2. NUEVOS CONTADORES (Esto es lo que faltaba)
    total_usuarios = Usuario.objects.count()
    total_eventos = Evento.objects.count()

    # Debug en consola para confirmar
    print(f"--- DEBUG HOME ADMIN ---")
    print(f"Usuarios: {total_usuarios}")
    print(f"Eventos: {total_eventos}")

    data = {
        'registros': registros,
        'Universidads': Universidads,
        # Pasamos los nuevos datos al template:
        'total_usuarios': total_usuarios,
        'total_eventos': total_eventos,
    }
    return render(request, 'home_admin.html', data)

def home_empresa(request, id):
    registros = RegistroActividad.objects.filter(usuario=request.user)
    
    todas_universidades = Universidad.objects.all()
    universidad_seleccionada = Universidad.objects.filter(id_universidad=id)

    data = {
        'registros': registros,  
        'empresa': universidad_seleccionada, 
        'empresas': todas_universidades
    }

    return render(request, 'empresa_1/home_empresa.html', data)


def tablasExtraccion(request,id):
        if request.user.is_authenticated:


                Universidads = Universidad.objects.all()
                Universidad = Universidad.objects.filter(id_universidad = id)

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapas = Etapa.objects.values_list("id_etapa", flat=True).filter(nombre="Extraccion materia prima")
                etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
                #Universidad = Universidad.objects.get(id)
                UniversidadArea = RegistroActividad.objects.all()
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)

                if request.user.is_staff:
                        entradas = Entrada.objects.filter(
                                etapa_id=etapa,
                                carrera__departamento__facultad__universidad__id_universidad=id  # filtra por la Universidad seleccionada
                        )
                else:
                        # Si no es admin, que solo vea sus propias entradas
                        entradas = Entrada.objects.filter(
                                usuario=request.user,
                                etapa_id=etapa
                        )


                #/////////////// Levenshtein ///////////////
                lista_u = []
                plabra_es = []
                lista_t = [] 
                nombre_espa = ""
                for e in entradas:
                        for a in area:
                                if e.carrera_id == a.carrera:        
                                        #print("las notas son: ",e.nombre) 
                                        e_nombre = e.nombre
                                        result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                                        #print(result)
                                        #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                                        if result > 1 :
                                                nombre_espa = e.nombre
                                                print(nombre_espa)
                                                separador = " "
                                                maximo_numero_de_separaciones = 2
                                                separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                                                plabra_es = plabra_es + separado_por_espacios
                                                
                                        else:
                                                lista_u.append(e_nombre)

                lista_t =  plabra_es + lista_u                       
                #print(lista_t)                    

                #////////////////////////////////////////
                
                result = (Entrada.objects
                .values('carrera')
                .annotate(dcount=Count('carrera'))
                .order_by()
                )
                #print("variable result!!!!!!!: ",result)

                theanswer = Entrada.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                #print(theanswer)
                salidas_count = Salida.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                
                oportunidad_count = Oportunidades.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)

                t = theanswer[0] if theanswer else None

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'area' : area,
                't': t,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'Universidads':Universidads,
                'Universidad':Universidad,
                'lista_t':lista_t,
                }
       
                return render(request,'Universidad_1/tablas_extraccion.html', data)
        else:
                return render(request, 'Universidad_1/tabla_extraccion.html')


def tablasDiseño(request,id):
        if request.user.is_authenticated:
                Universidads = Universidad.objects.all()
                Universidad = Universidad.objects.filter(id_universidad = id)

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Diseno y produccion"
                #Universidad = Universidad.objects.get(id)
                UniversidadArea = RegistroActividad.objects.all()
                
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                entradas = Entrada.objects.filter(etapa_id = etapa)
                
                
                result = (Entrada.objects
                .values('carrera')
                .annotate(dcount=Count('carrera'))
                .order_by()
                )
                print("variable result!!!!!!!: ",result)

               

                theanswer = Entrada.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                print(theanswer)
                salidas_count = Salida.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                
                oportunidad_count = Oportunidades.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
              
                
                                
                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'Universidads':Universidads,
                'Universidad':Universidad
                

                }
       
                return render(request,'Universidad_1/tablas_diseño.html', data)
        else:
                return render(request, 'Universidad_1/tabla_diseño.html')


def tablasLogistica(request,id):
        if request.user.is_authenticated:
                Universidads = Universidad.objects.all()
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
                UniversidadArea = RegistroActividad.objects.all()
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                theanswer = Entrada.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                salidas_count = Salida.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                oportunidad_count = Oportunidades.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                            
                data = {

                'registros': registros,
                'UniversidadArea':UniversidadArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'Universidads':Universidads,
                'Universidad':Universidad
        
                }
       
                return render(request,'Universidad_1/tablas_logistica.html', data)
        else:
                return render(request, 'Universidad_1/tabla_logistica.html')


def tablasCompra(request,id):
        if request.user.is_authenticated:
                Universidads = Universidad.objects.all()
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
                UniversidadArea = RegistroActividad.objects.all()
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                theanswer = Entrada.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                salidas_count = Salida.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                oportunidad_count = Oportunidades.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                            
                data = {

                'registros': registros,
                'UniversidadArea':UniversidadArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'Universidads':Universidads,
                'Universidad':Universidad
        
                }
       
                return render(request,'Universidad_1/tablas_compra.html', data)
        else:
                return render(request, 'Universidad_1/tabla_compra.html')

def tablasUso(request,id):
        if request.user.is_authenticated:
                Universidads = Universidad.objects.all()
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
                UniversidadArea = RegistroActividad.objects.all()
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                theanswer = Entrada.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                salidas_count = Salida.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                oportunidad_count = Oportunidades.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                            
                data = {

                'registros': registros,
                'UniversidadArea':UniversidadArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'Universidads':Universidads,
                'Universidad':Universidad
        
                }
       
                return render(request,'Universidad_1/tablas_uso.html', data)
        else:
                return render(request, 'Universidad_1/tabla_uso.html')


def tablasFin(request,id):
        if request.user.is_authenticated:
                Universidads = Universidad.objects.all()
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
                UniversidadArea = RegistroActividad.objects.all()
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                theanswer = Entrada.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                salidas_count = Salida.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                oportunidad_count = Oportunidades.objects.values('carrera').annotate(Count('carrera')).filter(etapa_id = etapa)
                            
                data = {

                'registros': registros,
                'UniversidadArea':UniversidadArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'Universidads':Universidads,
                'Universidad':Universidad
        
                }
       
                return render(request,'Universidad_1/tablas_fin.html', data)
        else:
                return render(request, 'Universidad_1/tabla_fin.html')



def promedioHome(request,id):
        if request.user.is_authenticated:
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'Universidad':Universidad,    

                }
       
                return render(request,"Universidad_1/promedios/promedio_home.html", data)
        else:
                return render(request, "Universidad_1/promedios/promedio_home.html")





def homeFrecuenciaDiseño(request,id):
        if request.user.is_authenticated:
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'Universidad':Universidad,    

                }   
                return render(request,"diseñoProduccion/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "Universidad_1/promedios/promedio_home.html")


def homeFrecuenciaLogistica(request,id):
        if request.user.is_authenticated:
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'Universidad':Universidad,    

                }   
                return render(request,"logistica/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "logistica/frecuencia/home_frecuencia.html")

def homeFrecuenciaCompra(request,id):
        if request.user.is_authenticated:
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'Universidad':Universidad,    

                }   
                return render(request,"compra/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "compra/frecuencia/home_frecuencia.html")





def homeFrecuenciaUso(request,id):
        if request.user.is_authenticated:
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'Universidad':Universidad,    

                }   

                return render(request,"usoConsumo/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "compra/frecuencia/home_frecuencia.html")



def homeFrecuenciaFin(request,id):
        if request.user.is_authenticated:
                Universidad = Universidad.objects.filter(id_universidad = id)
                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
                area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'Universidad':Universidad,    

                }   

                return render(request,"finVida/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "compra/frecuencia/home_frecuencia.html")




def promedioArea(request, id):
    # Usamos pk=id para asegurar que buscamos por la llave primaria de la carrera
    area = Carrera.objects.filter(pk=id)
    etapa = Etapa.objects.get(nombre="Extraccion materia prima")
    registros = RegistroActividad.objects.filter(usuario=request.user)
    
    b = 0
    Universidad_id = None # Inicializamos
    
    for a in area:
        if b < 1 :
            # CORRECCIÓN CRÍTICA: Navegamos la jerarquía completa
            # Carrera -> Departamento -> Facultad -> Universidad
            Universidad_id = a.departamento.facultad.universidad.id_universidad
            b = b + 1
            
    print("ID Universidad encontrada:", Universidad_id)
    
    # Evitamos error si no se encuentra universidad
    if Universidad_id:
        universidad_qs = Universidad.objects.filter(id_universidad=Universidad_id)
    else:
        universidad_qs = Universidad.objects.none()

    # Filtros de Entradas/Salidas/Oportunidades
    # Nota: Si tus modelos Entrada/Salida tienen el campo 'carrera', esto está bien:
    entradas = Entrada.objects.filter(carrera_id=id, etapa_id=etapa.id_etapa)
    salidas = Salida.objects.filter(carrera_id=id, etapa_id=etapa.id_etapa)
    oportunidades = Oportunidades.objects.filter(carrera_id=id, etapa_id=etapa.id_etapa)
    
    total_entradas = entradas.count()
    total_salidas = salidas.count()
    total_oportunidades = oportunidades.count()

    print("La id de la etapa es:", etapa)

    # /////////////// Entradas (Lógica de palabras) ///////////////
    lista_u = []
    plabra_es = []
    lista_t = [] 
    nombre_espa = ""

    for e in entradas:
        e_nombre = e.nombre
        result = len(e_nombre.split()) 
        
        if result > 1 :
            nombre_espa = e.nombre
            separador = " "
            maximo_numero_de_separaciones = 2
            separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
            plabra_es = plabra_es + separado_por_espacios
        else:
            lista_u.append(e_nombre)
            
    lista_t =  plabra_es + lista_u                          
    c = collections.Counter(lista_t)

    clave = c.keys()
    valor = c.values()
    cantidad_datos = c.items()

    # /////////////// Salidas ///////////////
    lista_unita_salida = []
    plabra_espacio_salida = []
    lista_total_salida = [] 
    nombre_espa_con_espacio = ""

    for e in salidas:
        e_nombre = e.nombre
        result = len(e_nombre.split())
        
        if result > 1 :
            nombre_espa_con_espacio = e.nombre
            separador = " "
            maximo_numero_de_separaciones = 2
            separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
            plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
        else:
            lista_unita_salida.append(e_nombre)

    lista_total_salida = plabra_espacio_salida + lista_unita_salida
    c_salidas = collections.Counter(lista_total_salida)

    clave_salidas = c_salidas.keys()
    valor_salidas = c_salidas.values()
    cantidad_datos_salidas = c_salidas.items()

    # /////////////// Oportunidades ///////////////
    lista_unita_opor = []
    plabra_espacio_opor = [] 
    lista_total_opor = [] 
    nombre_con_espacio_opor = "" 

    for e in oportunidades:
        e_nombre = e.nombre
        result = len(e_nombre.split())
        
        if result > 1 :
            nombre_con_espacio_opor = e.nombre
            separador = " "
            maximo_numero_de_separaciones = 2
            separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
            plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
        else:
            lista_unita_opor.append(e_nombre)

    lista_total_opor = plabra_espacio_opor + lista_unita_opor
    c_oportunidad = collections.Counter(lista_total_opor)

    clave_oportunidad = c_oportunidad.keys()
    valor_oportunidad = c_oportunidad.values()
    cantidad_datos_oportunidad = c_oportunidad.items()

    # /////////////// Estadísticas Generales ///////////////
    nota_masRepetida = ""
    b_val = 0
    for i in c:
        n = c[i]
        if b_val < n:
            b_val = n 
        if n == b_val :
            nota_masRepetida = i    

    # /////////////// Levenshtein (Matriz) ///////////////
    n_filas = len(lista_t)          
    n_columnas = len(lista_t)       
    
    # Validación por si lista_t está vacía para evitar errores
    if n_filas > 0:
        A = np.zeros([n_filas, n_columnas]) 
        for i in range(n_filas):
            for j in range(n_columnas):
                a_txt = lista_t[i]
                b_txt = lista_t[j]
                A[i,j] = round(jaro(a_txt, b_txt)*100)
        arr = A
    else:
        A = []
        arr = []

    # /////////////// Contexto Final ///////////////
    data = {
        'registros': registros,
        'area': area,
        'etapa': etapa,
        'Universidad': universidad_qs, # Usamos la variable corregida
        'total_entradas': total_entradas,
        'total_salidas': total_salidas,
        'total_oportunidades': total_oportunidades, # Corregí typo (estaba repetido total_salidas)
        'nota_masRepetida': nota_masRepetida,
        'lista_t': lista_t,

        # Entradas
        'clave': clave,
        'valor': valor,
        'cantidad_datos': cantidad_datos,

        # Salidas
        'clave_salidas': clave_salidas,
        'valor_salidas': valor_salidas,
        'cantidad_datos_salidas': cantidad_datos_salidas,

        # Oportunidad
        'clave_oportunidad': clave_oportunidad,
        'valor_oportunidad': valor_oportunidad,
        'cantidad_datos_oportunidad': cantidad_datos_oportunidad,

        'A': A,
        'arr': arr,
        'n_filas': n_filas,
        'n_columnas': n_columnas,
    }

    return render(request, 'Universidad_1/promedios/promedio.html', data)


def frecuenciaDiseño(request, id):
        area = Carrera.objects.filter(carrera = id)
        etapa = Etapa.objects.get(nombre = "Diseño y produccion")
        registros = RegistroActividad.objects.filter(usuario=request.user)    
        b = 0
        Universidad = 0
        for a in area:
                if b < 1 :
                        Universidad_id = a.id_universidad_id
                        b = b + 1
        print(Universidad_id)        

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id )


        
        entradas = Entrada.objects.filter(carrera_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(carrera_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(carrera_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(carrera_id = id).count()
        total_salidas = Salida.objects.filter(carrera_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(carrera_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'Universidad'                       :       Universidad,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'diseñoProduccion/frecuencia/tablas_frecuencia.html', data)   






def frecuenciaLogistica(request, id):
        area = Carrera.objects.filter(carrera = id)
        etapa = Etapa.objects.get(nombre = "Logistica")
        registros = RegistroActividad.objects.filter(usuario=request.user)    
        b = 0
        Universidad = 0
        for a in area:
                if b < 1 :
                        Universidad_id = a.id_universidad_id
                        b = b + 1
        print(Universidad_id)        

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id )


        
        entradas = Entrada.objects.filter(carrera_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(carrera_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(carrera_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(carrera_id = id).count()
        total_salidas = Salida.objects.filter(carrera_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(carrera_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'Universidad'                       :       Universidad,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'logistica/frecuencia/tablas_frecuencia.html', data)   



def frecuenciaCompra(request, id):
        area = Carrera.objects.filter(carrera = id)
        etapa = Etapa.objects.get(nombre = "Compra")
        registros = RegistroActividad.objects.filter(usuario=request.user)    
        b = 0
        Universidad = 0
        for a in area:
                if b < 1 :
                        Universidad_id = a.id_universidad_id
                        b = b + 1
        print(Universidad_id)        

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id )


        
        entradas = Entrada.objects.filter(carrera_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(carrera_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(carrera_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(carrera_id = id).count()
        total_salidas = Salida.objects.filter(carrera_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(carrera_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'Universidad'                       :       Universidad,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'compra/frecuencia/tablas_frecuencia.html', data)  



def frecuenciaUso(request, id):
        area = Carrera.objects.filter(carrera = id)
        etapa = Etapa.objects.get(nombre = "Uso consumo")
        registros = RegistroActividad.objects.filter(usuario=request.user)    
        b = 0
        Universidad = 0
        for a in area:
                if b < 1 :
                        Universidad_id = a.id_universidad_id
                        b = b + 1
        print(Universidad_id)        

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id )


        
        entradas = Entrada.objects.filter(carrera_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(carrera_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(carrera_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(carrera_id = id).count()
        total_salidas = Salida.objects.filter(carrera_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(carrera_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'Universidad'                       :       Universidad,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'usoConsumo/frecuencia/tablas_frecuencia.html', data)





def frecuenciaFin(request, id):
        area = Carrera.objects.filter(carrera = id)
        etapa = Etapa.objects.get(nombre = "Uso consumo")
        registros = RegistroActividad.objects.filter(usuario=request.user)    
        b = 0
        Universidad = 0
        for a in area:
                if b < 1 :
                        Universidad_id = a.id_universidad_id
                        b = b + 1
        print(Universidad_id)        

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id )


        
        entradas = Entrada.objects.filter(carrera_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(carrera_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(carrera_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(carrera_id = id).count()
        total_salidas = Salida.objects.filter(carrera_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(carrera_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'Universidad'                       :       Universidad,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'finVida/frecuencia/tablas_frecuencia.html', data)


# def promedioArea(request, id):
#         etapa = Etapa.objects.get(nombre = "Extraccion materia prima") 
#         registros = RegistroActividad.objects.filter(usuario=request.user)
#         Universidads = Universidad.objects.all()
#         Universidad = Universidad.objects.filter(id_universidad = id)
#         area =  Carrera.objects.filter(departamento__facultad__universidad__id_universidad= id)
#         UniversidadArea = RegistroActividad.objects.all()

#         entradas = Entrada.objects.filter(carrera_id = id, etapa_id = etapa)
#         salidas = Salida.objects.filter(carrera_id = id, etapa_id = etapa)
#         oportunidades = Oportunidades.objects.filter(carrera_id = id, etapa_id = etapa)
#         total_entradas = Entrada.objects.filter(carrera_id = id).count()
#         total_salidas = Salida.objects.filter(carrera_id = id).count()
#         total_oportunidades = Oportunidades.objects.filter(carrera_id = id).count()

#         print("la id de la estapa es!!!!!!!!: ",etapa)
        

# #/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
#         lista_u = []
#         plabra_es = []
#         lista_t = [] 
#         nombre_espa = ""

#         for e in entradas:
#                 #print("las notas son: ",e.nombre) 
#                 e_nombre = e.nombre
#                 result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
#                 #print(result)
#                 #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
#                 if result > 1 :
#                         nombre_espa = e.nombre
#                         #print(nombre_espa)
#                         separador = " "
#                         maximo_numero_de_separaciones = 2
#                         separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
#                         plabra_es = plabra_es + separado_por_espacios
                                                
#                 else:
#                         lista_u.append(e_nombre)
                        
#         lista_t =  plabra_es + lista_u                       
#         c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

#         clave = c.keys()
#         valor = c.values()
#         cantidad_datos = c.items()

#         # for clave, valor in cantidad_datos:
#         #         print (clave , ": " , valor)
 
# #//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

# #/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
#         lista_unita_salida = []
#         plabra_espacio_salida = []
#         lista_total_salida = [] 
#         nombre_espa_con_espacio = ""

#         for e in salidas:
#                 e_nombre = e.nombre
#                 result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
#                 # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
#                 if result > 1 :
#                         nombre_espa_con_espacio = e.nombre
                
#                         separador = " "
#                         maximo_numero_de_separaciones = 2
#                         separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
#                         plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
#                 else:
#                         lista_unita_salida.append(e_nombre)

#         lista_total_salida = plabra_espacio_salida + lista_unita_salida

#         print(lista_total_salida)
#         c_salidas = collections.Counter(lista_total_salida)
#         print(c_salidas)


#         clave_salidas = c_salidas.keys()
#         valor_salidas = c_salidas.values()
#         cantidad_datos_salidas = c_salidas.items()

#         # for clave, valor in cantidad_datos:
#         #         print (clave , ": " , valor)
 
# #//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

# #/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

#         lista_unita_opor = []                           #nota sin espacio en su nombre
#         plabra_espacio_opor = []                        #nota con espacio en su nombre
#         lista_total_opor = []                           #arrays con todas las notas
#         nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

#         for e in oportunidades:
                
#                 e_nombre = e.nombre
#                 result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
#                 # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
#                 if result > 1 :
#                         nombre_con_espacio_opor = e.nombre
                
#                         separador = " "
#                         maximo_numero_de_separaciones = 2
#                         separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
#                         plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
#                 else:
#                         lista_unita_opor.append(e_nombre)

#         lista_total_opor = plabra_espacio_opor + lista_unita_opor

#         # print(lista_total_salida)
#         c_oportunidad = collections.Counter(lista_total_opor)
#         print("las oportunidades son:  ", c_oportunidad)


#         clave_oportunidad = c_oportunidad.keys()
#         valor_oportunidad = c_oportunidad.values()
#         cantidad_datos_oportunidad = c_oportunidad.items()

#         # for clave_oportunidad, valor_oportunidad in cantidad_datos:
#         #         print (clave_oportunidad , ": " , valor_oportunidad)
 
# #//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 


#         values = c.values()
#         #print(c['harina'])
#         p = c['harina']
        

#         nota_masRepetida = ""
#         b = 0
#         for i in c:
#                 n = c[i]
#                 if b < n:
#                         b = n 
                               
#                 if n == b :
#                         nota_masRepetida = i
#         #print("la nota mas repetida es: ", nota_masRepetida) 
       

#         #///////////// Levenshetein ///////////////////////////
#         import numpy as np
#         # filas = len(lista_t)
#         # columnas = len(lista_t)


#         # A = np.zeros([filas,columnas])

#         # for i in range(filas):
#         #         for j in range(columnas):
#         #                 a = i+1
#         #                 x = j+1
#         #                 A[i,j] = a+x
#         # print(A)

#         n_filas = len(lista_t)          #Cantidad de registros
#         n_columnas = len(lista_t)       #Cantidad de registros
        
#         A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

#         for i in range(n_filas):
#                 for j in range(n_columnas):
#                         a = lista_t[i]
#                         b = lista_t[j]
#                         A[i,j] = round(jaro(a,b)*100)
                       
#         #print(A)   
#         # for line in A:
#         #         c = map
#         #         print (' '.join(map(str, line)))
#         arr = A
#         # for i in A:
#         #         print (i)
       
#         # for i in range(n_filas):
#         #         canti_registros =  i
#         #         columna = [fila[i] for fila in A] 

#                 #print(columna)    
#         #print(canti_registros)
#         #////////////////////////////////////////                

#         data = {

                
#                 'registros'                     :       registros,
#                 'Universidads'                      :       Universidads,
#                 'Universidad'                       :       Universidad,
#                 'total_entradas'                :       total_entradas,
#                 'total_salidas'                 :       total_salidas,
#                 'total_oportunidades'           :       total_salidas,
#                 'nota_masRepetida'              :       nota_masRepetida,
#                 'lista_t'                       :       lista_t,

#                 #entradas
#                 'clave'                         :       clave,
#                 'valor'                         :       valor,
#                 'cantidad_datos'                :       cantidad_datos,

#                 #salidas
#                 'clave_salidas'                 :       clave_salidas,
#                 'valor_salidas'                 :       valor_salidas,
#                 'cantidad_datos_salidas'        :       cantidad_datos_salidas,

#                 #oportunidad
#                 'clave_oportunidad'             :       clave_oportunidad,
#                 'valor_oportunidad'             :       valor_oportunidad,
#                 'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

#                 'A':A,
#                 'arr'                           :       arr,
#                 'n_filas'                       :       n_filas,
#                 'n_columnas'                    :       n_columnas,
#                 #'canti_registros'               :       canti_registros,
#                 'area'                          :       area,
#                 "UniversidadArea"                   :       UniversidadArea
 
#         }

#         return render(request, 'Universidad_1/promedios/promedio.html', data)


def entradasExtraccion(request):

    if request.user.is_authenticated:

        # =====================================================
        # 1. RECIBIR LA Universidad DESDE LA URL
        # =====================================================
        Universidad_id = request.GET.get("Universidad")
        Universidad_seleccionada = None
        Universidads = Universidad.objects.all()

        if Universidad_id:
            Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

        # =====================================================
        # 2. DATOS GENERALES
        # =====================================================
        registros = RegistroActividad.objects.filter(usuario=request.user)
        etapa = Etapa.objects.filter(nombre="Extraccion materia prima").first()

        if etapa:
            entradas = Entrada.objects.filter(etapa=etapa)
        else:
            entradas = Entrada.objects.none()

        UniversidadArea = RegistroActividad.objects.all()

        # =====================================================
        # 3. NUBE DE PALABRAS
        # =====================================================
        textos = " ".join(entradas.values_list("nombre", flat=True))
        nube_base64 = None

        if textos.strip():
            wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
            buffer = BytesIO()
            wc.to_image().save(buffer, format='PNG')
            nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

        # =====================================================
        # 4. ENVIAR TODO AL TEMPLATE (CORREGIDO)
        # =====================================================
        data = {
            'registros': registros,
            'entradas': entradas,
            'UniversidadArea': UniversidadArea,
            'Universidads': Universidads,
            'Universidad_seleccionada': Universidad_seleccionada,
            'nube_base64': nube_base64,
            'etapa_seleccionada': "extraccion",
        }

        return render(request, 'extraccion/entrada/tabla_entrada.html', data)

    else:
        return render(request, 'extraccion/entrada/tabla_entrada.html')


def SalidasExtraccion(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Extraccion materia prima").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)


                 # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'nube_base64': nube_base64,
                'etapa_seleccionada': "extraccion",

                }

       
                return render(request,'extraccion/salida/tabla_salida.html', data)
        else:
                return render(request, 'extraccion/salida/tabla_salida.html')
        

def OportunidadExtraccion(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Extraccion materia prima").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'nube_base64': nube_base64,
                'etapa_seleccionada': "extraccion",

                }
       
                return render(request,'extraccion/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'extraccion/oportunidad/tabla_oportunidad.html')



# Diseño y gestion     

def EntradaDiseño(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()


                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Diseño y produccion").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'nube_base64': nube_base64,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "diseñoProduccion",


                }

       
                return render(request,'diseñoProduccion/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'diseñoProduccion/entrada/tabla_entrada.html')

def salidaDiseño(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Diseño y produccion").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'nube_base64': nube_base64,
                'etapa_seleccionada': "diseñoProduccion",

                }

                return render(request,'diseñoProduccion/salida/tabla_salida.html', data)
        else:
                return render(request, 'diseñoProduccion/salida/tabla_salida.html')                              

def oportunidadDiseño(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Diseño y produccion").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)


                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'nube_base64': nube_base64,
                'etapa_seleccionada': "diseñoProduccion",

                }
       
                return render(request,'diseñoProduccion/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'diseñoProduccion/oportunidad/tabla_oportunidad.html')      


#logistica


def EntradaLogistica(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Logistica").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "logistica",
                'nube_base64': nube_base64,

                }

                return render(request,'logistica/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'logistica/entrada/tabla_entrada.html')


def salidaLogistica(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Logistica").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "logistica",
                'nube_base64': nube_base64,

                }
       
                return render(request,'logistica/salida/tabla_salida.html', data)
        else:
                return render(request, 'logistica/salida/tabla_salida.html') 


def oportunidadLogistica(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Logistica").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "logistica",
                'nube_base64': nube_base64,

                }
       
                return render(request,'logistica/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'logistica/oportunidad/tabla_oportunidad.html')                                     


# compra
def entradaCompra(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Compra").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "compra",
                'nube_base64': nube_base64,

                }
       
                return render(request,'compra/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'compra/entrada/tabla_entrada.html')      

def salidaCompra(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Compra").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "compra",
                'nube_base64': nube_base64,

                }

                return render(request,'compra/salida/tabla_salida.html', data)
        else:
                return render(request, 'compra/salida/tabla_salida.html') 


def oportunidadesCompra(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Compra").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "compra",
                'nube_base64': nube_base64,

                }
       
                return render(request,'compra/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'compra/salida/tabla_oportunidad.html') 

#Uso consumo

def entradaUsoConsumo(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Uso consumo").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "usoConsumo",
                'nube_base64': nube_base64,

                }
       
                return render(request,'usoConsumo/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'usoConsumo/entrada/tabla_entrada.html') 


def salidaUsoConsumo(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Uso consumo").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "usoConsumo",
                'nube_base64': nube_base64,

                }
       
                return render(request,'usoConsumo/salida/tabla_salida.html', data)
        else:
                return render(request, 'usoConsumo/salida/tabla_salida.html')


def oportunidadUsoConsumo(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Uso consumo").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "usoConsumo",
                'nube_base64': nube_base64,

                }
       
                return render(request,'usoConsumo/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'usoConsumo/oportunidad/tabla_oportunidad.html')



#Fin de vida
def entradaFin(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Fin de vida").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')


                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "finVida",
                'nube_base64': nube_base64,

                }
       
                return render(request,'finVida/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'finVida/entrada/tabla_entrada.html')               

def salidaFin(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Fin de vida").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "finVida",
                'nube_base64': nube_base64,

                }
       
                return render(request,'finVida/salida/tabla_salida.html', data)
        else:
                return render(request, 'finVida/salida/tabla_salida.html')

def oportunidadFin(request):

        if request.user.is_authenticated:

                Universidad_id = request.GET.get("Universidad")
                Universidad_seleccionada = None
                Universidads = Universidad.objects.all()

                if Universidad_id:
                        Universidad_seleccionada = Universidad.objects.filter(id_universidad=Universidad_id).first()

                registros = RegistroActividad.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Fin de vida").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                UniversidadArea = RegistroActividad.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'UniversidadArea':UniversidadArea,
                'Universidads': Universidads,
                'Universidad_seleccionada': Universidad_seleccionada,
                'etapa_seleccionada': "finVida",
                'nube_base64': nube_base64,

                }
       
                return render(request,'finVida/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'finVida/oportunidad/tabla_oportunidad.html')






# ############################################# graficos ###############################################

def homeGraficos(request):
        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        

        data = {
                'registros': registros,
                'Universidads': Universidads 
                }

        return render(request, 'graficos/home_graficos.html', data)



def etapaGraficos(request, id):
        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        area = Carrera.objects.filter(departamento__facultad__universidad__id_universidad= id)
        Universidad = Universidad.objects.filter(id_universidad = id)
        

        data = {
                'registros': registros,
                'Universidads': Universidads,
                'area':area,
                'Universidad':Universidad,
                }

        return render(request, 'graficos/etapas.html', data)



def areasExtraccion(request, id):
        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        Universidad = Universidad.objects.filter(id_universidad = id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
           


        data = {
                'registros': registros,
                'Universidads': Universidads,
                'Universidad':Universidad,
                'areas':areas
                }

        return render(request, 'extraccion/grafico/areas.html', data)




def graficosExtraccion(request, id):

        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        area_grafico = Carrera.objects.filter(carrera = id)
     
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, carrera = id)
        salidas = Salida.objects.filter(etapa_id = etapa, carrera = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, carrera = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        Universidad_id = i.id_universidad_id
                        b1 = b1 + 1


        Universidad = Universidad.objects.filter(id_universidad = Universidad_id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = Universidad_id)


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()


        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()


        ######## entradas ##############
        fechas = []
        for e in entradas:
                fechas.append(e.fecha.strftime("%d/%m/%Y"))

        c = collections.Counter(fechas)
           

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()    


        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'Universidads': Universidads,
                'area_grafico':area_grafico,
                'areas':areas,
                'Universidad':Universidad,
                'dias_total':dias_total,
                'clave': clave,
                'valor': valor,
                'cantidad_datos' : cantidad_datos,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'extraccion/grafico/graficos.html', data)       


def areasDiseño(request, id):
        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        Universidad = Universidad.objects.filter(id_universidad = id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
           


        data = {
                'registros': registros,
                'Universidads': Universidads,
                'Universidad':Universidad,
                'areas':areas
                }

        return render(request, 'diseñoProduccion/grafico/areas.html', data)




def graficosDiseño(request, id):

        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        area_grafico = Carrera.objects.filter(carrera = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, carrera = id)
        salidas = Salida.objects.filter(etapa_id = etapa, carrera = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, carrera = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        Universidad_id = i.id_universidad_id
                        b1 = b1 + 1

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = Universidad_id)
            


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items() 



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'Universidads': Universidads,
                'area_grafico':area_grafico,
                'areas':areas,
                'Universidad':Universidad,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'diseñoProduccion/grafico/graficos.html', data)       




def areasLogistica(request, id):
        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        Universidad = Universidad.objects.filter(id_universidad = id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
           


        data = {
                'registros': registros,
                'Universidads': Universidads,
                'Universidad':Universidad,
                'areas':areas
                }

        return render(request, 'logistica/grafico/areas.html', data)        




def graficosLogistica(request, id):

        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        area_grafico = Carrera.objects.filter(carrera = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, carrera = id)
        salidas = Salida.objects.filter(etapa_id = etapa, carrera = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, carrera = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        Universidad_id = i.id_universidad_id
                        b1 = b1 + 1

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = Universidad_id)
            


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'Universidads': Universidads,
                'area_grafico':area_grafico,
                'areas':areas,
                'Universidad':Universidad,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'logistica/grafico/graficos.html', data)         



def areasCompra(request, id):
        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        Universidad = Universidad.objects.filter(id_universidad = id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
           


        data = {
                'registros': registros,
                'Universidads': Universidads,
                'Universidad':Universidad,
                'areas':areas
                }

        return render(request, 'compra/grafico/areas.html', data)        




def graficosCompra(request, id):

        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        area_grafico = Carrera.objects.filter(carrera = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, carrera = id)
        salidas = Salida.objects.filter(etapa_id = etapa, carrera = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, carrera = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        Universidad_id = i.id_universidad_id
                        b1 = b1 + 1

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = Universidad_id)    


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'Universidads': Universidads,
                'area_grafico':area_grafico,
                'areas':areas,
                'Universidad':Universidad,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'compra/grafico/graficos.html', data)     



def areasUso(request, id):
        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        Universidad = Universidad.objects.filter(id_universidad = id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
           


        data = {
                'registros': registros,
                'Universidads': Universidads,
                'Universidad':Universidad,
                'areas':areas
                }

        return render(request, 'usoConsumo/grafico/areas.html', data)        



def graficosUso(request, id):

        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        area_grafico = Carrera.objects.filter(carrera = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del grafico la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, carrera = id)
        salidas = Salida.objects.filter(etapa_id = etapa, carrera = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, carrera = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        Universidad_id = i.id_universidad_id
                        b1 = b1 + 1

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = Universidad_id)  
            


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))
       

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'Universidads': Universidads,
                'area_grafico':area_grafico,
                'areas':areas,
                'Universidad':Universidad,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'usoConsumo/grafico/graficos.html', data)     




def areasFin(request, id):
        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        Universidad = Universidad.objects.filter(id_universidad = id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = id)
           


        data = {
                'registros': registros,
                'Universidads': Universidads,
                'Universidad':Universidad,
                'areas':areas
                }

        return render(request, 'finVida/grafico/areas.html', data)      


def graficosFin(request, id):

        registros = RegistroActividad.objects.filter(usuario = request.user)
        Universidads = Universidad.objects.all()
        area_grafico = Carrera.objects.filter(carrera = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, carrera = id)
        salidas = Salida.objects.filter(etapa_id = etapa, carrera = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, carrera = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        Universidad_id = i.id_universidad_id
                        b1 = b1 + 1

        Universidad = Universidad.objects.filter(id_universidad = Universidad_id)
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad = Universidad_id)  
            


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'Universidads': Universidads,
                'area_grafico':area_grafico,
                'areas':areas,
                'Universidad':Universidad,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'finVida/grafico/graficos.html', data)     





# reportes de excel

#  ///////////////////////////////////////// reportes de excel extraccion ///////////////////////////////////////////////////////////////////////////

class ReporteExcel(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 

        for e in entradas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalida(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidades(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response




#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////

#  ///////////////////////////////////////// reportes de excel Diseño y produccion ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaDiseño(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4

        for e in entradas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaDiseño(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadDiseño(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response




#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////


#  ///////////////////////////////////////// reportes de excel Logistica ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaLogistica(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 

        for e in entradas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaLogistica(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadLogistica(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response




#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////



#  ///////////////////////////////////////// reportes de excel Compra ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaCompra(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4

        for e in entradas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaCompra(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadCompra(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Oportunidades'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response




#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////


#  ///////////////////////////////////////// reportes de excel Uso Consumo ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaUso(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4

        for e in entradas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaUso(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadUso(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Oportunidades'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////


#  ///////////////////////////////////////// reportes de excel Fin de vida  ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaFin(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4

        for e in entradas:
                print(e.carrera.departamento.facultad.universidad.id_universidad)
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaFin(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadFin(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Oportunidades'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Universidad'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                ws.cell(row = cont, column = 2).value = e.carrera.departamento.facultad.universidad.id_universidad
                ws.cell(row = cont, column = 3).value = e.carrera_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response



def log_telegan(request):
        registros = RegistroActividad.objects.filter(usuario=request.user)
        logs_telegram = LogTelegram.objects.all()

        data = {
                'logs_telegram':logs_telegram,
                'registros': registros
        }
        return render(request,'log_telegram/log_telegram.html', data)

load_dotenv()

def leer_archivo(archivo):
    """
    Lee el contenido de un archivo subido (.txt, .docx, .pdf)
    y devuelve su texto como string.
    """
    extension = archivo.name.split('.')[-1].lower()

    try:
        if extension == "txt":
            return archivo.read().decode("utf-8", errors="ignore")

        elif extension == "docx":
            doc = docx.Document(archivo)
            return "\n".join([p.text for p in doc.paragraphs])

        elif extension == "pdf":
            texto = ""
            with fitz.open(stream=archivo.read(), filetype="pdf") as pdf:
                for pagina in pdf:
                    texto += pagina.get_text()
            return texto

    except Exception as e:
        return f"Error al leer el archivo: {str(e)}"

    return ""

def ia_semantica(request):
    """
    Vista principal del módulo de IA Semántica:
    - Permite subir un archivo o escribir texto libre
    - Genera un resumen de máximo 200 palabras usando OpenAI
    """
    resumen = None

    if request.method == 'POST':
        texto = request.POST.get('texto', '').strip()
        archivo = request.FILES.get('archivo')

        # Si se sube un archivo, leer su contenido
        if archivo:
            texto = leer_archivo(archivo)

            # Si el lector devolvió un mensaje de error o texto vacío
            if not texto or texto.startswith("Error"):
                resumen = texto if texto else "No se pudo leer el archivo o está vacío."
                return render(request, 'ia_semantica/ia_semantica.html', {'resumen': resumen})

        # Si hay texto, se envía a la API
        if texto:
            try:
                api_key = os.getenv("OPENAI_API_KEY")
                if not api_key:
                    resumen = "Error: No se encontró la clave OPENAI_API_KEY en el entorno."
                else:
                    client = OpenAI(api_key=api_key)

                    prompt = (
                        "Resume el siguiente texto en un máximo de 200 palabras, "
                        "manteniendo el tono formal y las ideas principales:\n\n" + texto
                    )

                    respuesta = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": "Eres un asistente experto en redacción de resúmenes en español."},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.7,
                        max_tokens=500
                    )

                    resumen = respuesta.choices[0].message.content.strip()


            except Exception as e:
                resumen = f"Error al procesar con la IA: {str(e)}"

        else:
            resumen = "Por favor, escribe un texto o sube un archivo antes de generar el resumen."

    return render(request, 'ia_semantica/ia_semantica.html', {'resumen': resumen})


# ======= Descargar resumen =======
def descargar_resumen(request):
    resumen = request.GET.get('resumen', '')

    if not resumen.strip():
        return HttpResponse("No hay resumen disponible para descargar.", content_type="text/plain")

    # Crear el documento Word
    buffer = BytesIO()
    doc = docx.Document()
    doc.add_heading("Resumen (200 palabras)", level=1)
    doc.add_paragraph(resumen)
    doc.save(buffer)
    buffer.seek(0)

    # Preparar la respuesta HTTP para descarga
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="resumen_200_palabras.docx"'
    return response


def procesamiento_area(request):
    Universidad_id = request.GET.get("Universidad")
    etapa_id = request.GET.get("etapa")
    entrada_key = request.GET.get("entrada")

    Universidads = Universidad.objects.all()
    Universidad_seleccionada = None
    etapa_seleccionada = None
    entrada_seleccionada = None
    etapa_obj = None
    texto_concatenado = ""
    resumen_ia = None

    # 1) Universidad
    if Universidad_id:
        try:
            Universidad_seleccionada = Universidad.objects.get(id_universidad=int(Universidad_id))
        except:
            Universidad_seleccionada = None

    # 2) Etapa
    if etapa_id:
        try:
            etapa_seleccionada = int(etapa_id)
            etapa_obj = Etapa.objects.get(id_etapa=etapa_seleccionada)
        except:
            etapa_seleccionada = None
            etapa_obj = None

    # 3) Tipo (entrada, salida, oportunidad)
    if entrada_key in ["entrada", "salida", "oportunidad"]:
        entrada_seleccionada = entrada_key

    # 4) Buscar información
    if Universidad_seleccionada and etapa_obj and entrada_seleccionada:

        if entrada_seleccionada == "entrada":
            registros = Entrada.objects.filter(
                etapa=etapa_obj,
                carrera__departamento__facultad__universidad__id_universidad=Universidad_seleccionada
            )

        elif entrada_seleccionada == "salida":
            registros = Salida.objects.filter(
                etapa=etapa_obj,
                carrera__departamento__facultad__universidad__id_universidad=Universidad_seleccionada
            )

        elif entrada_seleccionada == "oportunidad":
            registros = Oportunidades.objects.filter(
                etapa=etapa_obj,
                carrera__departamento__facultad__universidad__id_universidad=Universidad_seleccionada
            )

        texto_concatenado = " ".join([r.nombre for r in registros])

        # ----------- PROCESAR RESUMEN (POST) ----------------- 
        
        texto_a_resumir = request.POST.get("texto_concatenado", "")

        if texto_a_resumir:
            try:
                api_key = os.getenv("OPENAI_API_KEY")
                client = OpenAI(api_key=api_key)

                prompt = (
                    "Resume el siguiente texto en un máximo de 200 palabras, "
                    "manteniendo claridad, precisión y tono profesional:\n\n" +
                    texto_a_resumir
                )

                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Eres un experto en análisis y síntesis de texto."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=500,
                    temperature=0.5
                )

                resumen_ia = response.choices[0].message.content.strip()

            except Exception as e:
                resumen_ia = "Error al generar el resumen: " + str(e)

    return render(request, "procesamiento_area/procesamiento_area.html", {
        "Universidads": Universidads,
        "Universidad_seleccionada": Universidad_seleccionada,
        "etapa_seleccionada": etapa_seleccionada,
        "entrada_seleccionada": entrada_seleccionada,
        "texto_concatenado": texto_concatenado,
        "resumen_ia": resumen_ia,
    })


def procesamiento_ideas(request):
    Universidad_id = request.GET.get("Universidad")
    etapa_id = request.GET.get("etapa")

    Universidads = Universidad.objects.all()
    Universidad_seleccionada = None
    etapa_seleccionada = None
    etapa_obj = None
    texto_concatenado = ""
    resumen_ia = None

    # 1) Selección de Universidad
    if Universidad_id:
        try:
            Universidad_seleccionada = Universidad.objects.get(id_universidad=int(Universidad_id))
        except Universidad.DoesNotExist:
            Universidad_seleccionada = None

    # 2) Selección de etapa
    if etapa_id:
        try:
            etapa_seleccionada = int(etapa_id)
            etapa_obj = Etapa.objects.get(id_etapa=etapa_seleccionada)
        except Etapa.DoesNotExist:
            etapa_seleccionada = None
            etapa_obj = None

    # 3) Concatenar ideas de la Universidad y etapa
    if Universidad_seleccionada and etapa_obj:
        ideas = Idea.objects.filter(
            Universidad=Universidad_seleccionada,
            etapa=etapa_obj
        )
        texto_concatenado = " ".join([idea.texto for idea in ideas])

    # 4) Procesar resumen si se envía POST
    if request.method == "POST":
        texto_a_resumir = request.POST.get("texto_concatenado", "")
        if texto_a_resumir:
            try:
                api_key = os.getenv("OPENAI_API_KEY")
                client = OpenAI(api_key=api_key)

                prompt = (
                    "Resume el siguiente texto en un máximo de 200 palabras, "
                    "manteniendo claridad, precisión y tono profesional:\n\n" +
                    texto_a_resumir
                )

                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Eres un experto en análisis y síntesis de texto."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=500,
                    temperature=0.5
                )

                resumen_ia = response.choices[0].message.content.strip()

            except Exception as e:
                resumen_ia = "Error al generar el resumen: " + str(e)

    return render(request, "procesamiento_ideas/procesamiento_ideas.html", {
        "Universidads": Universidads,
        "Universidad_seleccionada": Universidad_seleccionada,
        "etapa_seleccionada": etapa_seleccionada,
        "etapa_obj": etapa_obj,
        "texto_concatenado": texto_concatenado,
        "resumen_ia": resumen_ia,
    })

def home_procesamiento(request):
    return render(request, "procesamiento/home_procesamiento.html")

def procesamiento_aempresa(request):
    Universidad_id = request.GET.get("Universidad")
    area_id = request.GET.get("area")

    Universidads = Universidad.objects.all()
    Universidad_seleccionada = None
    areas = None
    area_seleccionada = None

    texto_concatenado = ""
    resumen_ia = None

    # 1. Selección Universidad
    if Universidad_id:
        try:
            Universidad_seleccionada = Universidad.objects.get(id_universidad=int(Universidad_id))
        except Universidad.DoesNotExist:
            Universidad_seleccionada = None

    # 2. Cargar áreas de la Universidad seleccionada
    if Universidad_seleccionada:
        areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad=Universidad_seleccionada)

    # 3. Selección área
    if area_id:
        try:
            area_seleccionada = Carrera.objects.get(carrera=int(area_id))
        except Carrera.DoesNotExist:
            area_seleccionada = None

    # 4. Concatenación de entradas, salidas, oportunidades
    if area_seleccionada:
        entradas = Entrada.objects.filter(carrera=area_seleccionada)
        salidas = Salida.objects.filter(carrera=area_seleccionada)
        oportunidades = Oportunidades.objects.filter(carrera=area_seleccionada)

        texto_concatenado = " ".join(
            [e.nombre for e in entradas] +
            [s.nombre for s in salidas] +
            [o.nombre for o in oportunidades]
        )

    # 5. PROCESO DE RESUMEN IA
    if request.method == "POST":
        texto_a_resumir = request.POST.get("texto_concatenado", "")
        if texto_a_resumir:
            try:
                client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

                prompt = (
                    "Resume el siguiente texto en un máximo de 200 palabras, "
                    "manteniendo claridad, precisión y tono profesional:\n\n" +
                    texto_a_resumir
                )

                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Eres un experto en análisis y síntesis de texto."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=500,
                    temperature=0.5
                )

                resumen_ia = response.choices[0].message.content.strip()

            except Exception as e:
                resumen_ia = "Error al generar el resumen: " + str(e)

    return render(request, "procesamiento_aUniversidad/procesamiento_aUniversidad.html", {
        "Universidads": Universidads,
        "Universidad_seleccionada": Universidad_seleccionada,
        "areas": areas,
        "area_seleccionada": area_seleccionada,
        "texto_concatenado": texto_concatenado,
        "resumen_ia": resumen_ia,
    })


def generar_clave(longitud=10):
    caracteres = string.ascii_letters + string.digits
    return ''.join(random.choice(caracteres) for _ in range(longitud))

# -----------------------------------------------------------------------------
# 1. ADMIN USUARIOS (Listar)
# -----------------------------------------------------------------------------


@login_required
@user_passes_test(lambda u: u.is_staff)
def admin_usuarios(request):
    # 1. Obtener todas las universidades para los botones
    universidades = Universidad.objects.all()

    # Obtener ID de la URL
    uni_id = request.GET.get("universidad")

    universidad_seleccionada = None
    usuarios_universidad = None
    coordinador_universidad = None 

    if uni_id:
        try:
            universidad_seleccionada = Universidad.objects.get(pk=uni_id)

            # Consultamos DIRECTAMENTE a Usuario
            # (Asegúrate de que 'carrera' y la relación existen en tu modelo Usuario)
            usuarios_lista = Usuario.objects.filter(
                carrera__departamento__facultad__universidad=universidad_seleccionada
            ).select_related("carrera").order_by('last_name')

            from django.core.paginator import Paginator
            paginator = Paginator(usuarios_lista, 10)
            page_number = request.GET.get('page')
            usuarios_universidad = paginator.get_page(page_number)

            # C. Obtener Coordinador
            if hasattr(Usuario, 'universidad_coordinador'):
                coordinador_universidad = Usuario.objects.filter(
                    universidad_coordinador=universidad_seleccionada
                ).first()

        except Universidad.DoesNotExist:
            universidad_seleccionada = None

    context = {
        "universidades": universidades,
        "universidad_seleccionada": universidad_seleccionada,
        "usuarios_universidad": usuarios_universidad,
        "is_paginated": usuarios_universidad.has_other_pages() if usuarios_universidad else False,
        "coordinador_universidad": coordinador_universidad,
        
        # --- VARIABLES OBLIGATORIAS AGREGADAS ---
        # 1. Define qué plantilla extender. Si está en app/templates/base_admin.html,
        # Django la encuentra solo con el nombre:
        "base_template": "base_admin.html", 
        
        # 2. Nombre exacto de la URL en tu urls.py para el botón 'Volver'
        "volver_url": "home_admin", 
        
        # 3. Bandera para ocultar/mostrar elementos exclusivos de coordinador
        "es_coordinador_view": False, 
    }
    
    return render(request, "admin_usuarios/admin_usuarios.html", context)


# -----------------------------------------------------------------------------
# 2. CREAR USUARIO
# -----------------------------------------------------------------------------

@login_required
@user_passes_test(lambda u: u.is_staff)
def crear_usuario(request):
    print(f"MÉTODO: {request.method}, AJAX: {request.headers.get('x-requested-with')}, ACTION: {request.POST.get('action')}")
    uni_id = request.GET.get("universidad") or request.POST.get("universidad")
    try:
        uni_id = int(uni_id)
    except (TypeError, ValueError):
        uni_id = None
    
    # --- BLOQUE AJAX (API PROPIA + Deptos/Carreras) ---
    if request.POST.get('action'):
        action = request.POST.get('action')
        
        if action == 'buscar_deptos':
            fac_id = request.POST.get('id')
            deptos = Departamento.objects.filter(facultad_id=fac_id).values('pk', 'nombre').order_by('nombre')
            return JsonResponse([{'id': d['pk'], 'nombre': d['nombre']} for d in deptos], safe=False)
            
        elif action == 'buscar_carreras':
            depto_id = request.POST.get('id')
            carreras = Carrera.objects.filter(departamento_id=depto_id).values('pk', 'nombre').order_by('nombre')
            return JsonResponse([{'id': c['pk'], 'nombre': c['nombre']} for c in carreras], safe=False)

        # --- NUEVA API DE CHILE (Rápida y local) ---
        elif action == 'buscar_comunas':
            region_nombre = request.POST.get('id') # El JS manda el nombre de la región como ID
            print(f"🌍 Buscando comunas para: {region_nombre}")
            
            # Diccionario completo de Chile
            chile_data = {
                "Arica y Parinacota": ["Arica", "Camarones", "Putre", "General Lagos"],
                "Tarapacá": ["Iquique", "Alto Hospicio", "Pozo Almonte", "Camiña", "Colchane", "Huara", "Pica"],
                "Antofagasta": ["Antofagasta", "Mejillones", "Sierra Gorda", "Taltal", "Calama", "Ollagüe", "San Pedro de Atacama", "Tocopilla", "María Elena"],
                "Atacama": ["Copiapó", "Caldera", "Tierra Amarilla", "Chañaral", "Diego de Almagro", "Vallenar", "Alto del Carmen", "Freirina", "Huasco"],
                "Coquimbo": ["La Serena", "Coquimbo", "Andacollo", "La Higuera", "Paiguano", "Vicuña", "Illapel", "Canela", "Los Vilos", "Salamanca", "Ovalle", "Combarbalá", "Monte Patria", "Punitaqui", "Río Hurtado"],
                "Valparaíso": ["Valparaíso", "Casablanca", "Concón", "Juan Fernández", "Puchuncaví", "Quintero", "Viña del Mar", "Isla de Pascua", "Los Andes", "Calle Larga", "Rinconada", "San Esteban", "La Ligua", "Cabildo", "Papudo", "Petorca", "Zapallar", "Quillota", "Calera", "Hijuelas", "La Cruz", "Nogales", "San Antonio", "Algarrobo", "Cartagena", "El Quisco", "El Tabo", "Santo Domingo", "San Felipe", "Catemu", "Llaillay", "Panquehue", "Putaendo", "Santa María", "Quilpué", "Limache", "Olmué", "Villa Alemana"],
                "Metropolitana": ["Cerrillos", "Cerro Navia", "Conchalí", "El Bosque", "Estación Central", "Huechuraba", "Independencia", "La Cisterna", "La Florida", "La Granja", "La Pintana", "La Reina", "Las Condes", "Lo Barnechea", "Lo Espejo", "Lo Prado", "Macul", "Maipú", "Ñuñoa", "Pedro Aguirre Cerda", "Peñalolén", "Providencia", "Pudahuel", "Quilicura", "Quinta Normal", "Recoleta", "Renca", "San Joaquín", "San Miguel", "San Ramón", "Santiago", "Vitacura", "Puente Alto", "Pirque", "San José de Maipo", "Colina", "Lampa", "Tiltil", "San Bernardo", "Buin", "Calera de Tango", "Paine", "Melipilla", "Alhué", "Curacaví", "María Pinto", "San Pedro", "Talagante", "El Monte", "Isla de Maipo", "Padre Hurtado", "Peñaflor"],
                "O'Higgins": ["Rancagua", "Codegua", "Coinco", "Coltauco", "Doñihue", "Graneros", "Las Cabras", "Machalí", "Malloa", "Mostazal", "Olivar", "Peumo", "Pichidegua", "Quinta de Tilcoco", "Rengo", "Requínoa", "San Vicente", "Pichilemu", "La Estrella", "Litueche", "Marchihue", "Navidad", "Paredones", "San Fernando", "Chépica", "Chimbarongo", "Lolol", "Nancagua", "Palmilla", "Peralillo", "Placilla", "Pumanque", "Santa Cruz"],
                "Maule": ["Talca", "Constitución", "Curepto", "Empedrado", "Maule", "Pelarco", "Pencahue", "Río Claro", "San Clemente", "San Rafael", "Cauquenes", "Chanco", "Pelluhue", "Curicó", "Hualañé", "Licantén", "Molina", "Rauco", "Romeral", "Sagrada Familia", "Teno", "Vichuquén", "Linares", "Colbún", "Longaví", "Parral", "Retiro", "San Javier", "Villa Alegre", "Yerbas Buenas"],
                "Ñuble": ["Chillán", "Bulnes", "Cobquecura", "Coelemu", "Coihueco", "Chillán Viejo", "El Carmen", "Ninhue", "Ñiquén", "Pemuco", "Pinto", "Portezuelo", "Quillón", "Quirihue", "Ránquil", "San Carlos", "San Fabián", "San Ignacio", "San Nicolás", "Treguaco", "Yungay"],
                "Biobío": ["Concepción", "Coronel", "Chiguayante", "Florida", "Hualqui", "Lota", "Penco", "San Pedro de la Paz", "Santa Juana", "Talcahuano", "Tomé", "Hualpén", "Lebu", "Arauco", "Cañete", "Contulmo", "Curanilahue", "Los Álamos", "Tirúa", "Los Ángeles", "Antuco", "Cabrero", "Laja", "Mulchén", "Nacimiento", "Negrete", "Quilaco", "Quilleco", "San Rosendo", "Santa Bárbara", "Tucapel", "Yumbel", "Alto Biobío"],
                "Araucanía": ["Temuco", "Carahue", "Cunco", "Curarrehue", "Freire", "Galvarino", "Gorbea", "Lautaro", "Loncoche", "Melipeuco", "Nueva Imperial", "Padre Las Casas", "Perquenco", "Pitrufquén", "Pucón", "Saavedra", "Teodoro Schmidt", "Toltén", "Vilcún", "Villarrica", "Cholchol", "Angol", "Collipulli", "Curacautín", "Ercilla", "Lonquimay", "Los Sauces", "Lumaco", "Purén", "Renaico", "Traiguén", "Victoria"],
                "Los Ríos": ["Valdivia", "Corral", "Lanco", "Los Lagos", "Máfil", "Mariquina", "Paillaco", "Panguipulli", "La Unión", "Futrono", "Lago Ranco", "Río Bueno"],
                "Los Lagos": ["Puerto Montt", "Calbuco", "Cochamó", "Fresia", "Frutillar", "Los Muermos", "Llanquihue", "Maullín", "Puerto Varas", "Castro", "Ancud", "Chonchi", "Curaco de Vélez", "Dalcahue", "Puqueldón", "Queilén", "Quellón", "Quemchi", "Quinchao", "Osorno", "Puerto Octay", "Purranque", "Puyehue", "Río Negro", "San Juan de la Costa", "San Pablo", "Chaitén", "Futaleufú", "Hualaihué", "Palena"],
                "Aysén": ["Coyhaique", "Lago Verde", "Aysén", "Cisnes", "Guaitecas", "Cochrane", "O'Higgins", "Tortel", "Chile Chico", "Río Ibáñez"],
                "Magallanes": ["Punta Arenas", "Laguna Blanca", "Río Verde", "San Gregorio", "Cabo de Hornos (Williams)", "Antártica", "Porvenir", "Primavera", "Timaukel", "Natales", "Torres del Paine"]
            }
            
            comunas = chile_data.get(region_nombre, [])
            # Devolvemos el nombre como ID también, ya que guardamos texto en la BD
            return JsonResponse([{'id': c, 'nombre': c} for c in comunas], safe=False)

    # --- BLOQUE POST (Guardar Usuario) ---
    if request.method == "POST":
        username = request.POST.get('username')
        email = request.POST.get('email')
        nombre = request.POST.get('nombre')
        apellido = request.POST.get('apellido')
        telefono = request.POST.get('telefono')
        carrera_id = request.POST.get('carrera')
        
        # --- NUEVOS CAMPOS ---
        region = request.POST.get('region')
        comuna = request.POST.get('comuna')
        genero = request.POST.get('genero')
        fecha_nacimiento = request.POST.get('fecha_nacimiento')
        # ---------------------
        
        # Validaciones básicas
        if not username or not email or not nombre or not apellido:
            messages.error(request, "Todos los campos de texto son obligatorios.")
            return redirect(request.path + f"?universidad={uni_id}")

        if Usuario.objects.filter(username=username).exists():
            messages.error(request, "Ese Username ya existe.")
            return redirect(request.path + f"?universidad={uni_id}")

        try:
            clave_temporal = get_random_string(length=8)
            
            # Crear usuario base
            nuevo_usuario = Usuario.objects.create_user(
                username=username,
                email=email,
                password=clave_temporal,
                first_name=nombre,
                last_name=apellido
            )
            
            # --- GUARDAR TODOS LOS DATOS EXTRA ---
            # Usamos getattr/setattr o asignación directa si el modelo Usuario es personalizado
            if telefono: nuevo_usuario.telefono = telefono
            if region: nuevo_usuario.region = region
            if comuna: nuevo_usuario.comuna = comuna
            if genero: nuevo_usuario.genero = genero
            if fecha_nacimiento: nuevo_usuario.fecha_nacimiento = fecha_nacimiento
            # -------------------------------------
            
            # Asignar Carrera
            if carrera_id:
                try:
                    carrera_obj = Carrera.objects.get(pk=carrera_id)
                    nuevo_usuario.carrera = carrera_obj
                except Carrera.DoesNotExist:
                    pass 
            
            nuevo_usuario.save() # Guardamos todo junto

            # Enviar Correo (Tu lógica original intacta)
            try:
                asunto = "Bienvenido a Alumni - Activa tu cuenta"
                uid = urlsafe_base64_encode(force_bytes(nuevo_usuario.pk))
                token = default_token_generator.make_token(nuevo_usuario)
                domain = request.get_host()
                scheme = request.scheme
                reset_url = reverse('password_reset_confirm', kwargs={'uidb64': uid, 'token': token})
                full_url = f"{scheme}://{domain}{reset_url}"
                print(f"🔗 Link generado: {full_url}")
                mensaje = f"Hola {nombre},\n\nTu cuenta ha sido creada en la plataforma Alumni.\n\nUsuario: {username}\n\nPara activar tu cuenta y crear tu contraseña, haz clic aquí:\n{full_url}\n\nEste enlace expirará en 24 horas.\n\nBienvenido/a,\nEquipo Alumni"
                EmailThread(asunto, mensaje, [email]).start()
                print("✅ Correo enviado en segundo plano")
                messages.success(request, "Usuario creado exitosamente. Se envió el correo de activación.")
            except Exception as e:
                print(f"ERROR CORREO: {e}")
                messages.warning(request, f"Usuario creado, pero falló el correo: {e}")

            return redirect(f"/administrador/usuarios?universidad={uni_id}")

        except Exception as e:
            messages.error(request, f"Error al crear: {e}")
            return redirect(request.path + f"?universidad={uni_id}")

    # --- BLOQUE GET (Carga inicial) ---
    facultades = Facultad.objects.none()

    if uni_id is not None:
        facultades = Facultad.objects.filter(universidad_id=uni_id)

    return render(request, "coordinador/crear_usuario.html", {
        "facultades": facultades, 
        "universidad_id": uni_id
    })
@login_required
@user_passes_test(lambda u: u.is_staff)
def api_comunas(request):
    region_nombre = request.GET.get('region', '')
    chile_data = {
        "Arica y Parinacota": ["Arica", "Camarones", "Putre", "General Lagos"],
        "Tarapacá": ["Iquique", "Alto Hospicio", "Pozo Almonte", "Camiña", "Colchane", "Huara", "Pica"],
        "Antofagasta": ["Antofagasta", "Mejillones", "Sierra Gorda", "Taltal", "Calama", "Ollagüe", "San Pedro de Atacama", "Tocopilla", "María Elena"],
        "Atacama": ["Copiapó", "Caldera", "Tierra Amarilla", "Chañaral", "Diego de Almagro", "Vallenar", "Alto del Carmen", "Freirina", "Huasco"],
        "Coquimbo": ["La Serena", "Coquimbo", "Andacollo", "La Higuera", "Paiguano", "Vicuña", "Illapel", "Canela", "Los Vilos", "Salamanca", "Ovalle", "Combarbalá", "Monte Patria", "Punitaqui", "Río Hurtado"],
        "Valparaíso": ["Valparaíso", "Casablanca", "Concón", "Juan Fernández", "Puchuncaví", "Quintero", "Viña del Mar", "Isla de Pascua", "Los Andes", "Calle Larga", "Rinconada", "San Esteban", "La Ligua", "Cabildo", "Papudo", "Petorca", "Zapallar", "Quillota", "Calera", "Hijuelas", "La Cruz", "Nogales", "San Antonio", "Algarrobo", "Cartagena", "El Quisco", "El Tabo", "Santo Domingo", "San Felipe", "Catemu", "Llaillay", "Panquehue", "Putaendo", "Santa María", "Quilpué", "Limache", "Olmué", "Villa Alemana"],
        "Metropolitana": ["Cerrillos", "Cerro Navia", "Conchalí", "El Bosque", "Estación Central", "Huechuraba", "Independencia", "La Cisterna", "La Florida", "La Granja", "La Pintana", "La Reina", "Las Condes", "Lo Barnechea", "Lo Espejo", "Lo Prado", "Macul", "Maipú", "Ñuñoa", "Pedro Aguirre Cerda", "Peñalolén", "Providencia", "Pudahuel", "Quilicura", "Quinta Normal", "Recoleta", "Renca", "San Joaquín", "San Miguel", "San Ramón", "Santiago", "Vitacura", "Puente Alto", "Pirque", "San José de Maipo", "Colina", "Lampa", "Tiltil", "San Bernardo", "Buin", "Calera de Tango", "Paine", "Melipilla", "Alhué", "Curacaví", "María Pinto", "San Pedro", "Talagante", "El Monte", "Isla de Maipo", "Padre Hurtado", "Peñaflor"],
        "O'Higgins": ["Rancagua", "Codegua", "Coinco", "Coltauco", "Doñihue", "Graneros", "Las Cabras", "Machalí", "Malloa", "Mostazal", "Olivar", "Peumo", "Pichidegua", "Quinta de Tilcoco", "Rengo", "Requínoa", "San Vicente", "Pichilemu", "La Estrella", "Litueche", "Marchihue", "Navidad", "Paredones", "San Fernando", "Chépica", "Chimbarongo", "Lolol", "Nancagua", "Palmilla", "Peralillo", "Placilla", "Pumanque", "Santa Cruz"],
        "Maule": ["Talca", "Constitución", "Curepto", "Empedrado", "Maule", "Pelarco", "Pencahue", "Río Claro", "San Clemente", "San Rafael", "Cauquenes", "Chanco", "Pelluhue", "Curicó", "Hualañé", "Licantén", "Molina", "Rauco", "Romeral", "Sagrada Familia", "Teno", "Vichuquén", "Linares", "Colbún", "Longaví", "Parral", "Retiro", "San Javier", "Villa Alegre", "Yerbas Buenas"],
        "Ñuble": ["Chillán", "Bulnes", "Cobquecura", "Coelemu", "Coihueco", "Chillán Viejo", "El Carmen", "Ninhue", "Ñiquén", "Pemuco", "Pinto", "Portezuelo", "Quillón", "Quirihue", "Ránquil", "San Carlos", "San Fabián", "San Ignacio", "San Nicolás", "Treguaco", "Yungay"],
        "Biobío": ["Concepción", "Coronel", "Chiguayante", "Florida", "Hualqui", "Lota", "Penco", "San Pedro de la Paz", "Santa Juana", "Talcahuano", "Tomé", "Hualpén", "Lebu", "Arauco", "Cañete", "Contulmo", "Curanilahue", "Los Álamos", "Tirúa", "Los Ángeles", "Antuco", "Cabrero", "Laja", "Mulchén", "Nacimiento", "Negrete", "Quilaco", "Quilleco", "San Rosendo", "Santa Bárbara", "Tucapel", "Yumbel", "Alto Biobío"],
        "Araucanía": ["Temuco", "Carahue", "Cunco", "Curarrehue", "Freire", "Galvarino", "Gorbea", "Lautaro", "Loncoche", "Melipeuco", "Nueva Imperial", "Padre Las Casas", "Perquenco", "Pitrufquén", "Pucón", "Saavedra", "Teodoro Schmidt", "Toltén", "Vilcún", "Villarrica", "Cholchol", "Angol", "Collipulli", "Curacautín", "Ercilla", "Lonquimay", "Los Sauces", "Lumaco", "Purén", "Renaico", "Traiguén", "Victoria"],
        "Los Ríos": ["Valdivia", "Corral", "Lanco", "Los Lagos", "Máfil", "Mariquina", "Paillaco", "Panguipulli", "La Unión", "Futrono", "Lago Ranco", "Río Bueno"],
        "Los Lagos": ["Puerto Montt", "Calbuco", "Cochamó", "Fresia", "Frutillar", "Los Muermos", "Llanquihue", "Maullín", "Puerto Varas", "Castro", "Ancud", "Chonchi", "Curaco de Vélez", "Dalcahue", "Puqueldón", "Queilén", "Quellón", "Quemchi", "Quinchao", "Osorno", "Puerto Octay", "Purranque", "Puyehue", "Río Negro", "San Juan de la Costa", "San Pablo", "Chaitén", "Futaleufú", "Hualaihué", "Palena"],
        "Aysén": ["Coyhaique", "Lago Verde", "Aysén", "Cisnes", "Guaitecas", "Cochrane", "O'Higgins", "Tortel", "Chile Chico", "Río Ibáñez"],
        "Magallanes": ["Punta Arenas", "Laguna Blanca", "Río Verde", "San Gregorio", "Cabo de Hornos (Williams)", "Antártica", "Porvenir", "Primavera", "Timaukel", "Natales", "Torres del Paine"]
    }
    comunas = chile_data.get(region_nombre, [])
    return JsonResponse([{'id': c, 'nombre': c} for c in comunas], safe=False)

# -----------------------------------------------------------------------------
# 2. EDITAR USUARIO 
# -----------------------------------------------------------------------------
@login_required
@user_passes_test(lambda u: u.is_staff)
def editar_usuario(request, user_id):
    usuario = get_object_or_404(Usuario, id=user_id)
    
    # --- API DE COMUNAS AJAX PARA EDICIÓN ---
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        action = request.POST.get('action')
        print(f"AJAX recibido - action: {action}")
        if action == 'buscar_comunas':
            region_nombre = request.POST.get('id')
            chile_data = {
                "Arica y Parinacota": ["Arica", "Camarones", "Putre", "General Lagos"],
                "Tarapacá": ["Iquique", "Alto Hospicio", "Pozo Almonte", "Camiña", "Colchane", "Huara", "Pica"],
                "Antofagasta": ["Antofagasta", "Mejillones", "Sierra Gorda", "Taltal", "Calama", "Ollagüe", "San Pedro de Atacama", "Tocopilla", "María Elena"],
                "Atacama": ["Copiapó", "Caldera", "Tierra Amarilla", "Chañaral", "Diego de Almagro", "Vallenar", "Alto del Carmen", "Freirina", "Huasco"],
                "Coquimbo": ["La Serena", "Coquimbo", "Andacollo", "La Higuera", "Paiguano", "Vicuña", "Illapel", "Canela", "Los Vilos", "Salamanca", "Ovalle", "Combarbalá", "Monte Patria", "Punitaqui", "Río Hurtado"],
                "Valparaíso": ["Valparaíso", "Casablanca", "Concón", "Juan Fernández", "Puchuncaví", "Quintero", "Viña del Mar", "Isla de Pascua", "Los Andes", "Calle Larga", "Rinconada", "San Esteban", "La Ligua", "Cabildo", "Papudo", "Petorca", "Zapallar", "Quillota", "Calera", "Hijuelas", "La Cruz", "Nogales", "San Antonio", "Algarrobo", "Cartagena", "El Quisco", "El Tabo", "Santo Domingo", "San Felipe", "Catemu", "Llaillay", "Panquehue", "Putaendo", "Santa María", "Quilpué", "Limache", "Olmué", "Villa Alemana"],
                "Metropolitana": ["Cerrillos", "Cerro Navia", "Conchalí", "El Bosque", "Estación Central", "Huechuraba", "Independencia", "La Cisterna", "La Florida", "La Granja", "La Pintana", "La Reina", "Las Condes", "Lo Barnechea", "Lo Espejo", "Lo Prado", "Macul", "Maipú", "Ñuñoa", "Pedro Aguirre Cerda", "Peñalolén", "Providencia", "Pudahuel", "Quilicura", "Quinta Normal", "Recoleta", "Renca", "San Joaquín", "San Miguel", "San Ramón", "Santiago", "Vitacura", "Puente Alto", "Pirque", "San José de Maipo", "Colina", "Lampa", "Tiltil", "San Bernardo", "Buin", "Calera de Tango", "Paine", "Melipilla", "Alhué", "Curacaví", "María Pinto", "San Pedro", "Talagante", "El Monte", "Isla de Maipo", "Padre Hurtado", "Peñaflor"],
                "O'Higgins": ["Rancagua", "Codegua", "Coinco", "Coltauco", "Doñihue", "Graneros", "Las Cabras", "Machalí", "Malloa", "Mostazal", "Olivar", "Peumo", "Pichidegua", "Quinta de Tilcoco", "Rengo", "Requínoa", "San Vicente", "Pichilemu", "La Estrella", "Litueche", "Marchihue", "Navidad", "Paredones", "San Fernando", "Chépica", "Chimbarongo", "Lolol", "Nancagua", "Palmilla", "Peralillo", "Placilla", "Pumanque", "Santa Cruz"],
                "Maule": ["Talca", "Constitución", "Curepto", "Empedrado", "Maule", "Pelarco", "Pencahue", "Río Claro", "San Clemente", "San Rafael", "Cauquenes", "Chanco", "Pelluhue", "Curicó", "Hualañé", "Licantén", "Molina", "Rauco", "Romeral", "Sagrada Familia", "Teno", "Vichuquén", "Linares", "Colbún", "Longaví", "Parral", "Retiro", "San Javier", "Villa Alegre", "Yerbas Buenas"],
                "Ñuble": ["Chillán", "Bulnes", "Cobquecura", "Coelemu", "Coihueco", "Chillán Viejo", "El Carmen", "Ninhue", "Ñiquén", "Pemuco", "Pinto", "Portezuelo", "Quillón", "Quirihue", "Ránquil", "San Carlos", "San Fabián", "San Ignacio", "San Nicolás", "Treguaco", "Yungay"],
                "Biobío": ["Concepción", "Coronel", "Chiguayante", "Florida", "Hualqui", "Lota", "Penco", "San Pedro de la Paz", "Santa Juana", "Talcahuano", "Tomé", "Hualpén", "Lebu", "Arauco", "Cañete", "Contulmo", "Curanilahue", "Los Álamos", "Tirúa", "Los Ángeles", "Antuco", "Cabrero", "Laja", "Mulchén", "Nacimiento", "Negrete", "Quilaco", "Quilleco", "San Rosendo", "Santa Bárbara", "Tucapel", "Yumbel", "Alto Biobío"],
                "Araucanía": ["Temuco", "Carahue", "Cunco", "Curarrehue", "Freire", "Galvarino", "Gorbea", "Lautaro", "Loncoche", "Melipeuco", "Nueva Imperial", "Padre Las Casas", "Perquenco", "Pitrufquén", "Pucón", "Saavedra", "Teodoro Schmidt", "Toltén", "Vilcún", "Villarrica", "Cholchol", "Angol", "Collipulli", "Curacautín", "Ercilla", "Lonquimay", "Los Sauces", "Lumaco", "Purén", "Renaico", "Traiguén", "Victoria"],
                "Los Ríos": ["Valdivia", "Corral", "Lanco", "Los Lagos", "Máfil", "Mariquina", "Paillaco", "Panguipulli", "La Unión", "Futrono", "Lago Ranco", "Río Bueno"],
                "Los Lagos": ["Puerto Montt", "Calbuco", "Cochamó", "Fresia", "Frutillar", "Los Muermos", "Llanquihue", "Maullín", "Puerto Varas", "Castro", "Ancud", "Chonchi", "Curaco de Vélez", "Dalcahue", "Puqueldón", "Queilén", "Quellón", "Quemchi", "Quinchao", "Osorno", "Puerto Octay", "Purranque", "Puyehue", "Río Negro", "San Juan de la Costa", "San Pablo", "Chaitén", "Futaleufú", "Hualaihué", "Palena"],
                "Aysén": ["Coyhaique", "Lago Verde", "Aysén", "Cisnes", "Guaitecas", "Cochrane", "O'Higgins", "Tortel", "Chile Chico", "Río Ibáñez"],
                "Magallanes": ["Punta Arenas", "Laguna Blanca", "Río Verde", "San Gregorio", "Cabo de Hornos (Williams)", "Antártica", "Porvenir", "Primavera", "Timaukel", "Natales", "Torres del Paine"]
            }
            comunas = chile_data.get(region_nombre, [])
            return JsonResponse([{'id': c, 'nombre': c} for c in comunas], safe=False)

    # Determinar ID Universidad
    uni_id = None
    try:
        if usuario.universidad:
            uni_id = usuario.universidad.pk
        elif usuario.carrera:
            uni_id = usuario.carrera.departamento.facultad.universidad.pk
        else:
            uni_id = request.GET.get("universidad")
    except Exception:
        uni_id = request.GET.get("universidad")

    # Cargar carreras disponibles
    carreras = []
    if uni_id:
        carreras = Carrera.objects.filter(departamento__facultad__universidad__id_universidad=uni_id)

    if request.method == "POST":
        nombre = request.POST.get("nombre", "").strip()
        apellido = request.POST.get("apellido", "").strip()
        email = request.POST.get("email", "").strip()
        telefono = request.POST.get("telefono", "").strip()
        
        # --- CAMPOS NUEVOS ---
        region = request.POST.get("region")
        comuna = request.POST.get("comuna", "").strip()
        genero = request.POST.get("genero")
        fecha_nac = request.POST.get("fecha_nacimiento")
        # ---------------------
        
        carrera_id = request.POST.get("carrera")

        if not nombre:
            messages.error(request, "El nombre no puede estar vacío.")
            return redirect(request.path)

        # Actualizar Datos
        usuario.first_name = nombre
        usuario.last_name = apellido
        usuario.email = email
        if hasattr(usuario, 'telefono'):
            usuario.telefono = telefono if telefono else None
            
        usuario.region = region
        usuario.comuna = comuna
        usuario.genero = genero
        if fecha_nac: usuario.fecha_nacimiento = fecha_nac
        
        # Actualizar Carrera
        if carrera_id:
            try:
                nueva_carrera = Carrera.objects.get(pk=carrera_id)
                usuario.carrera = nueva_carrera
                
                # Actualizar legacy si existe
                RegistroActividad.objects.update_or_create(
                    usuario=usuario,
                    defaults={'carrera': nueva_carrera, 'descripcion': 'Actualizado por admin'}
                )
            except Carrera.DoesNotExist:
                messages.error(request, "La carrera seleccionada no existe.")
        
        usuario.save()
        messages.success(request, "Usuario actualizado correctamente.")
        
        if uni_id:
            return redirect(f"/administrador/usuarios?universidad={uni_id}")
        return redirect("admin_usuarios")

    return render(request, "admin_usuarios/editar_usuario.html", {
        "usuario": usuario,
        "carreras": carreras,
        "universidad_id": uni_id,
        "carrera_actual_id": usuario.carrera.pk if usuario.carrera else None
    })

# -----------------------------------------------------------------------------
# 3. ELIMINAR USUARIO 
# -----------------------------------------------------------------------------
@login_required
@user_passes_test(lambda u: u.is_staff)
def eliminar_usuario(request, id):
    usuario = get_object_or_404(Usuario, id=id)
    
    # Guardamos la U antes de borrar para saber dónde volver
    uni_id = None
    if usuario.carrera:
        uni_id = usuario.carrera.departamento.facultad.universidad.pk
    else:
        # Si no tiene carrera, intentamos buscar en registro antiguo
        reg = RegistroActividad.objects.filter(usuario=usuario).first()
        if reg: uni_id = reg.carrera.departamento.facultad.universidad.pk

    usuario.delete()
    messages.success(request, "Usuario eliminado correctamente.")
    
    if uni_id:
        return redirect(f"/administrador/usuarios?universidad={uni_id}")
    return redirect("admin_usuarios") # Ruta genérica si no sabemos la U


# -----------------------------------------------------------------------------
# 4. RESETEAR CLAVE (Corregido con Try-Except)
# -----------------------------------------------------------------------------
@login_required
@user_passes_test(lambda u: u.is_staff)
def resetear_clave(request, user_id):
    if request.method == "POST":
        usuario = get_object_or_404(Usuario, pk=user_id)
        
        # 1. Generar token de seguridad y codificar el ID del usuario
        uid = urlsafe_base64_encode(force_bytes(usuario.pk))
        token = default_token_generator.make_token(usuario)
        
        # 2. Construir la URL absoluta para el link
        domain = request.get_host()
        scheme = request.scheme
        reset_url = reverse('password_reset_confirm', kwargs={'uidb64': uid, 'token': token})
        full_url = f"{scheme}://{domain}{reset_url}"

        # ENVÍO BLINDADO
        try:
            send_mail(
                subject="Restablecimiento de contraseña - Alumni",
                message=f"Hola {usuario.first_name},\n\nUn administrador ha solicitado el restablecimiento de tu contraseña.\n\nHaz clic en el siguiente enlace para crear una nueva clave de forma segura:\n{full_url}\n\nEste enlace expirará pronto.\n\nSi no esperabas este correo, por favor avísanos.",
                from_email=settings.EMAIL_HOST_USER,
                recipient_list=[usuario.email],
                fail_silently=False,
            )
            messages.success(request, f"Enlace de recuperación enviado exitosamente a {usuario.email}")
        except Exception as e:
            print(f"⚠️ Error SMTP Reset: {e}")
            messages.error(request, f"Error al enviar el correo a {usuario.email}. Revisa la configuración de email.")

    # Redirección inteligente (Mantenida exactamente como la tenías)
    usuario_obj = get_object_or_404(Usuario, pk=user_id)
    if usuario_obj.carrera:
        uni_id = usuario_obj.carrera.departamento.facultad.universidad.pk
        return redirect(f"/administrador/usuarios?universidad={uni_id}")
        
    return redirect("admin_usuarios")


# Función auxiliar para generar clave (si no la tienes en otro lado)
def generar_clave(longitud=8):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=longitud))

def obtener_usuarios_sin_actividad(id_universidad):
    # todas las áreas de esa Universidad
    areas_ids = Carrera.objects.filter(departamento__facultad__universidad__id_universidad=id_universidad).values_list('carrera', flat=True)

    # usuarios asociados a esas áreas
    usuarios_ids = RegistroActividad.objects.filter(
        carrera__in=areas_ids
    ).values_list("usuario_id", flat=True)

    # usuarios sin actividad (no registraron NADA)
    usuarios_sin_act = Usuario.objects.filter(id__in=usuarios_ids)\
        .exclude(id__in=Entrada.objects.values_list("usuario_id", flat=True))\
        .exclude(id__in=Salida.objects.values_list("usuario_id", flat=True))\
        .exclude(id__in=Oportunidades.objects.values_list("usuario_id", flat=True))

    return usuarios_sin_act

def notificaciones(request):
    Universidads = Carrera.objects.values(
        "id_universidad__id_universidad",
        "id_universidad__nombre"
    ).distinct()

    Universidad_id = request.GET.get("Universidad")
    usuarios_faltantes = []
    total_usuarios = []

    if Universidad_id:
        usuarios_faltantes = obtener_usuarios_sin_actividad(Universidad_id)

        total_usuarios = Usuario.objects.filter(
            RegistroActividad__carrera__departamento__facultad__universidad__id_universidad=Universidad_id
        ).distinct().count()

    context = {
        "Universidads": Universidads,
        "Universidad_id": Universidad_id,
        "usuarios_faltantes": usuarios_faltantes,
        "total_usuarios": total_usuarios,
    }

    return render(request, "notificaciones/notificaciones.html", context)

def enviar_recordatorio_view(request, id_universidad):

    # 1) Obtener usuarios sin actividad
    usuarios = obtener_usuarios_sin_actividad(id_universidad)

    # 2) Si no faltan usuarios, volver con mensaje
    if not usuarios.exists():
        messages.warning(
            request,
            "Todos los usuarios ya han respondido. No hay faltantes."
        )
        return redirect(f"/administrador/notificaciones/?Universidad={id_universidad}")

    enviados = 0

    # 3) Enviar correo a cada usuario
    for u in usuarios:
        if not u.email:
            continue

        try:
            send_mail(
                subject="Recordatorio — Lineal a Circular",
                message=(
                    f"Hola {u.first_name or u.username},\n\n"
                    "Te recordamos completar tus respuestas de Entradas, "
                    "Salidas y Oportunidades.\n\n"
                    "Gracias,\nEquipo Ciclo Circular."
                ),
                from_email=settings.EMAIL_HOST_USER,
                recipient_list=[u.email],  # correo individual
                fail_silently=False,
            )
            enviados += 1
        except Exception as e:
            print(f"Error enviando a {u.email}: {e}")

    # 4) Mensaje de éxito
    messages.success(request, f"Emails enviados correctamente: {enviados}")

    # 5) Volver a notificaciones con la misma Universidad filtrada
    return redirect(f"/administrador/notificaciones/?Universidad={id_universidad}")


def enviar_mensaje_todos(request, id_universidad):
    if request.method == "POST":
        mensaje = request.POST.get("mensaje")
        archivo = request.FILES.get("archivo")

        # Obtener usuarios de la Universidad (por los registros)
        usuarios = Usuario.objects.filter(RegistroActividad__carrera__departamento__facultad__universidad__id_universidad=id_universidad).distinct()

        enviados = 0

        for u in usuarios:
            if u.email:

                email = EmailMessage(
                    subject="Mensaje de la Universidad",
                    body=mensaje,
                    from_email=settings.EMAIL_HOST_USER,
                    to=[u.email],
                )

                # Si hay archivo adjunto
                if archivo:
                    email.attach(archivo.name, archivo.read(), archivo.content_type)

                try:
                    email.send()
                    enviados += 1
                except Exception as e:
                    print("Error enviando a:", u.email, e)

        messages.success(request, f"Mensaje enviado correctamente a {enviados} usuarios.")
        return redirect(f"/administrador/notificaciones/?Universidad={id_universidad}")

def form_mensaje_todos(request, id_universidad):
    return render(request, "notificaciones/form_mensaje_todos.html", {
        "id_universidad": id_universidad
    })

def menu_notificaciones(request):
    return render(request, "notificaciones/menu_notificaciones.html")

def notificaciones_comuna(request):

    # Obtener SOLO comunas de usuarios NO admin
    comunas = (
        RegistroActividad.objects
        .filter(usuario__is_staff=False)
        .values_list("carrera__comuna", flat=True)
        .distinct()
    )

    comuna_seleccionada = request.GET.get("comuna")

    total_usuarios = 0

    if comuna_seleccionada:
        total_usuarios = (
            Usuario.objects
            .filter(RegistroActividad__carrera__comuna=comuna_seleccionada,
                    is_staff=False)
            .distinct()
            .count()
        )

    context = {
        "comunas": comunas,
        "comuna_seleccionada": comuna_seleccionada,
        "total_usuarios": total_usuarios,
    }

    return render(request, "notificaciones/notificaciones_comuna.html", context)


def form_mensaje_comuna(request, comuna):
    return render(request, "notificaciones/form_mensaje_comuna.html", {
        "comuna": comuna
    })

def enviar_mensaje_comuna(request, comuna):
    if request.method == "POST":
        mensaje = request.POST.get("mensaje")
        archivo = request.FILES.get("archivo")

        # Filtrar usuarios por comuna
        usuarios = Usuario.objects.filter(
            RegistroActividad__carrera__comuna=comuna,
            is_staff=False
        ).distinct()

        enviados = 0

        for u in usuarios:
            if not u.email:
                continue

            email = EmailMessage(
                subject=f"Mensaje para usuarios de {comuna}",
                body=mensaje,
                from_email=settings.EMAIL_HOST_USER,
                to=[u.email],
            )

            if archivo:
                email.attach(archivo.name, archivo.read(), archivo.content_type)

            try:
                email.send()
                enviados += 1
            except Exception as e:
                print("Error:", u.email, e)

        messages.success(request, f"Mensaje enviado a {enviados} usuarios de {comuna}.")
        return redirect("/administrador/notificaciones/comuna/")


def ver_cv(request, user_id):
    usuario = get_object_or_404(Usuario, id=user_id)
    cv = usuario.cvusuario_set.first()

    palabras = []
    if cv:
        palabras = [
            cv.palabra1, cv.palabra2, cv.palabra3, cv.palabra4, cv.palabra5,
            cv.palabra6, cv.palabra7, cv.palabra8, cv.palabra9, cv.palabra10,
        ]
        # Filtramos los None o cadenas vacías
        palabras = [p for p in palabras if p]

    return render(request, 'coordinador/ver_cv.html', {
        'usuario': usuario,
        'cv': cv,
        'palabras': palabras,
    })

def agregar_cv(request, user_id):
    usuario = get_object_or_404(Usuario, id=user_id)

    if request.method == "POST":
        archivo = request.FILES.get("cv")

        if archivo:
            archivo_bytes = archivo.read()
            archivo_base64 = base64.b64encode(archivo_bytes).decode("utf-8")

            CVUsuario.objects.create(
                usuario=usuario,
                nombre_archivo=archivo.name,
                archivo=archivo_base64
            )

            return redirect("ver_cv", user_id=usuario.id)

    return render(request, "coordinador/agregar_cv.html", {
        "usuario": usuario
    })
    
def descargar_cv(request, cv_id):
    cv = get_object_or_404(CVUsuario, id=cv_id)
    # Convertir Base64 a bytes si está almacenado así
    try:
        archivo_bytes = base64.b64decode(cv.archivo)
    except Exception:
        archivo_bytes = cv.archivo  # si está en bytes puro

    response = HttpResponse(archivo_bytes, content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{cv.nombre_archivo}"'
    return response

def procesamiento_palabra_clave(request):
    Universidads = Universidad.objects.all()

    Universidad_id = request.GET.get("Universidad")
    Universidad_seleccionada = None
    usuarios = []
    matriz = []
    matriz_con_usuarios = []

    if Universidad_id:
        Universidad_seleccionada = Universidad.objects.get(id_universidad=Universidad_id)

        # Usuarios de la Universidad con CV
        usuarios = Usuario.objects.filter(
            RegistroActividad__carrera__departamento__facultad__universidad__id_universidad__id_universidad=Universidad_id,
            cvusuario__isnull=False
        ).distinct()

        if request.method == "POST":

            # Extraer palabras clave por usuario
            lista_palabras = []
            for usr in usuarios:
                cv = CVUsuario.objects.filter(usuario=usr).first()

                if cv:
                    palabras = [
                        cv.palabra1, cv.palabra2, cv.palabra3, cv.palabra4, cv.palabra5,
                        cv.palabra6, cv.palabra7, cv.palabra8, cv.palabra9, cv.palabra10
                    ]
                    palabras = [p.strip().lower() for p in palabras if p]
                else:
                    palabras = []

                lista_palabras.append(palabras)

            tamaño = len(usuarios)

            # Crear matriz vacía NxN
            matriz = [[0] * tamaño for _ in range(tamaño)]

            # Calcular coincidencias
            for i in range(tamaño):
                for j in range(tamaño):

                    # Diagonal siempre vale 10
                    if i == j:
                        matriz[i][j] = 10
                        continue

                    coincidencias = len(set(lista_palabras[i]) & set(lista_palabras[j]))
                    matriz[i][j] = coincidencias

            # Emparejar usuarios con su fila de la matriz
            matriz_con_usuarios = list(zip(usuarios, matriz))

    return render(request, 'procesamiento_palabra_clave/procesamiento_palabra_clave.html', {
        "Universidads": Universidads,
        "Universidad_seleccionada": Universidad_seleccionada,
        "usuarios": usuarios,
        "matriz_con_usuarios": matriz_con_usuarios,
    })


def es_super_admin(user):
    return user.is_staff


@user_passes_test(lambda u: u.is_superuser)
def asignar_coordinador(request, universidad_id):
    if request.method == "POST":
        try:
            user_id = request.POST.get("user_id")

            usuario = get_object_or_404(Usuario, id=user_id)
            universidad_obj = get_object_or_404(Universidad, pk=universidad_id)

            # 1. Quitar al coordinador anterior (usando minúscula 'universidad_coordinador')
            Usuario.objects.filter(
                universidad_coordinador=universidad_obj
            ).update(
                universidad_coordinador=None,
                es_coordinador=False
            )

            # 2. Asignar el nuevo (usando minúscula 'universidad_coordinador')
            usuario.universidad_coordinador = universidad_obj
            usuario.es_coordinador = True
            usuario.save()

            messages.success(request, f"Se ha asignado a {usuario.username} como coordinador correctamente.")

        except Exception as e:
            messages.error(request, f"Error al asignar coordinador: {e}")

    return redirect(request.META.get("HTTP_REFERER", "/"))


# ============================================
# QUITAR COORDINADOR (CORREGIDO)
# ============================================
@user_passes_test(lambda u: u.is_superuser)
def quitar_coordinador(request, user_id):
    try:
        usuario = get_object_or_404(Usuario, id=user_id)

        # Usando minúscula 'universidad_coordinador'
        usuario.universidad_coordinador = None
        usuario.es_coordinador = False
        usuario.save()

        messages.success(request, f"Se ha quitado el rol de coordinador a {usuario.username}.")
        
    except Exception as e:
        messages.error(request, f"Error al quitar coordinador: {e}")

    return redirect(request.META.get('HTTP_REFERER', '/'))

# ============================================
# PANEL DEL COORDINADOR
# ============================================
@login_required
def panel_coordinador(request):
    usuario = request.user
    if not usuario.es_coordinador or not usuario.universidad_coordinador:
        return redirect('home') 

    from app.models import Evento
    eventos = Evento.objects.filter(
        universidad=usuario.universidad_coordinador
    ).order_by('-inicio')

    return render(request, "administrador/panel_coordinador.html", {
        "universidad": usuario.universidad_coordinador,
        "eventos": eventos,
    })


# ============================================
# REGISTRO DEL COORDINADOR
# ============================================
@login_required
def registro(request):
    # Si no es coordinador → fuera
    if not request.user.Universidad_coordinador:
        return redirect('home')

    Universidad = request.user.Universidad_coordinador

    etapa_slug = request.GET.get('etapa')
    tabla = request.GET.get('tabla')  # entrada, salida, oportunidad

    MAPA_ETAPAS = {
        'extraccion': 1,
        'diseno': 2,
        'logistica': 3,
        'compra': 4,
        'uso': 5,
        'fin': 6,
    }

    id_etapa = MAPA_ETAPAS.get(etapa_slug)
    etapa_obj = Etapa.objects.filter(id_etapa=id_etapa).first()

    registros = []

    # ----------------------------
    # Selección dinámica de tabla
    # ----------------------------
    MODELOS = {
        "entrada": Entrada,
        "salida": Salida,
        "oportunidad": Oportunidades,
    }

    Modelo = MODELOS.get(tabla)

    if Modelo and etapa_obj:
        registros = Modelo.objects.filter(
            carrera__departamento__facultad__universidad__id_universidad=Universidad,
            etapa=etapa_obj
        )

    return render(request, 'coordinador/registro.html', {
        'Universidad': Universidad,
        'etapa': etapa_slug,
        'tabla': tabla,
        'registros': registros,
    })

def frecuencias(request):
    # -------------------------
    #  Mapa de slugs a ID de etapa (local dentro del def)
    # -------------------------
    MAPA_ETAPAS = {
        'extraccion': 1,
        'diseno': 2,
        'logistica': 3,
        'compra': 4,
        'uso': 5,
        'fin': 6,
    }

    # -------------------------
    #  Validar que el coordinador tenga Universidad
    # -------------------------
    Universidad = request.user.Universidad_coordinador
    if not Universidad:
        return render(request, "error.html", {"mensaje": "No tienes una Universidad asignada."})

    # -------------------------
    #  Obtener etapa desde GET
    # -------------------------
    etapa_slug = request.GET.get("etapa")
    if not etapa_slug:
        # Página inicial sin etapa seleccionada
        return render(request, "coordinador/frecuencias.html", {
            "etapa": None,
            "areas": None,
            "entradas_count": [],
            "salidas_count": [],
            "oportunidades_count": [],
        })

    # -------------------------
    #  Obtener el objeto Etapa
    # -------------------------
    id_etapa = MAPA_ETAPAS.get(etapa_slug)
    if not id_etapa:
        return render(request, "error.html", {"mensaje": "Etapa inválida."})

    try:
        etapa_obj = Etapa.objects.get(id_etapa=id_etapa)
    except Etapa.DoesNotExist:
        return render(request, "error.html", {"mensaje": "La etapa seleccionada no existe."})

    # -------------------------
    # Obtener áreas de la Universidad
    # -------------------------
    areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad=Universidad.id_universidad)

    # -------------------------
    #  Contar Entradas / Salidas / Oportunidades por área
    # -------------------------
    entradas_count = (
        Entrada.objects.filter(etapa=etapa_obj, carrera__in=areas)
        .values("carrera")
        .annotate(total=Count("carrera"))
    )

    salidas_count = (
        Salida.objects.filter(etapa=etapa_obj, carrera__in=areas)
        .values("carrera")
        .annotate(total=Count("carrera"))
    )

    oportunidades_count = (
        Oportunidades.objects.filter(etapa=etapa_obj, carrera__in=areas)
        .values("carrera")
        .annotate(total=Count("carrera"))
    )

    # -------------------------
    #  Renderizar template
    # -------------------------
    return render(request, "coordinador/frecuencias.html", {
        "etapa": etapa_slug,
        "etapa_display": etapa_obj.nombre,
        "areas": areas,
        "entradas_count": entradas_count,
        "salidas_count": salidas_count,
        "oportunidades_count": oportunidades_count,
    })

def graficos(request):
    # -----------------------
    # 1. Validar coordinador
    # -----------------------
    Universidad = getattr(request.user, 'Universidad_coordinador', None)
    if not Universidad:
        return render(request, "error.html", {"mensaje": "No tienes una Universidad asignada."})

    # -----------------------
    # 2. Diccionario de etapas
    # -----------------------
    etapas = {
        "extraccion": "Extracción Materia Prima",
        "diseno": "Diseño y Producción",
        "logistica": "Logística",
        "compra": "Compra",
        "uso": "Uso y Consumo",
        "fin": "Fin de Vida",
    }

    # -----------------------
    # 3. Obtener etapa seleccionada
    # -----------------------
    etapa_slug = request.GET.get("etapa")
    etapa_obj = None
    if etapa_slug:
        # Buscar objeto Etapa
        for slug, nombre in etapas.items():
            if slug == etapa_slug:
                etapa_obj = Etapa.objects.filter(nombre=nombre).first()
                break

    # -----------------------
    # 4. Obtener áreas de la Universidad
    # -----------------------
    areas = Carrera.objects.filter(departamento__facultad__universidad__id_universidad=Universidad.id_universidad)

    # -----------------------
    # 5. Obtener área seleccionada
    # -----------------------
    area_id = request.GET.get("area")
    area_grafico = None
    cantidad_datos_dicc = []
    cantidad_datos_dicc_salida = []
    cantidad_datos_dicc_oportunidad = []

    if area_id and etapa_obj:
        area_grafico = areas.filter(carrera=area_id)

        # Entradas
        entradas = (
            Entrada.objects.filter(etapa=etapa_obj, carrera__carrera=area_id)
            .values('fecha')
            .annotate(total=Count('carrera'))
            .order_by('fecha')
        )
        cantidad_datos_dicc = [(str(e['fecha']), e['total']) for e in entradas]

        # Salidas
        salidas = (
            Salida.objects.filter(etapa=etapa_obj, carrera__carrera=area_id)
            .values('fecha')
            .annotate(total=Count('carrera'))
            .order_by('fecha')
        )
        cantidad_datos_dicc_salida = [(str(s['fecha']), s['total']) for s in salidas]

        # Oportunidades
        oportunidades = (
            Oportunidades.objects.filter(etapa=etapa_obj, carrera__carrera=area_id)
            .values('fecha')
            .annotate(total=Count('carrera'))
            .order_by('fecha')
        )
        cantidad_datos_dicc_oportunidad = [(str(o['fecha']), o['total']) for o in oportunidades]

    # -----------------------
    # 6. Renderizar template
    # -----------------------
    return render(request, "coordinador/graficos.html", {
        "etapas": etapas,
        "etapa": etapa_slug,
        "areas": areas,
        "area_grafico": area_grafico,
        "cantidad_datos_dicc": cantidad_datos_dicc,
        "cantidad_datos_dicc_salida": cantidad_datos_dicc_salida,
        "cantidad_datos_dicc_oportunidad": cantidad_datos_dicc_oportunidad,
    })

@login_required
def coordina_usuarios(request):
    try:
        usuario = request.user
        
        # 1. Seguridad
        if not usuario.es_coordinador or not usuario.universidad_coordinador:
            return redirect('home')

        Universidad_seleccionada = usuario.universidad_coordinador
        
        # 2. Filtrado
        usuarios = Usuario.objects.filter(
            carrera__departamento__facultad__universidad=Universidad_seleccionada,
            is_active=True
        ).select_related('carrera').order_by('last_name')

        # 3. Contexto
        context = {
            "usuarios_universidad": usuarios,
            "universidad_seleccionada": Universidad_seleccionada,
            "coordinador_universidad": usuario,
            "es_coordinador_view": True,
            
            # Navegación dinámica
            "base_template": "base_public.html",
            "volver_url": "panel_coordinador",
            
            "universidades": [] 
        }

        # --- CORRECCIÓN DEFINITIVA: Ruta exacta según tu estructura ---
        return render(request, "admin_usuarios/admin_usuarios.html", context)

    except Exception as e:
        print(f"🔴 ERROR FATAL EN COORDINA_USUARIOS: {str(e)}")
        return HttpResponse(f"Error interno: {str(e)}", status=500)
    

@login_required
def crear_usuario_coordinacion(request):
    usuario = request.user
    
    # 1. Seguridad: Verificar rol y asignación
    if not getattr(usuario, 'es_coordinador', False) or not getattr(usuario, 'universidad_coordinador', None):
        return redirect('home')

    universidad_obj = usuario.universidad_coordinador

    # ---------------------------------------------------------
    # 2. BLOQUE AJAX (Para los selectores en cascada)
    # ---------------------------------------------------------
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        action = request.POST.get('action')
        
        if action == 'buscar_deptos':
            fac_id = request.POST.get('id')
            # Validamos que la facultad pertenezca a la U del coordinador por seguridad
            deptos = Departamento.objects.filter(
                facultad_id=fac_id, 
                facultad__universidad=universidad_obj
            ).values('pk', 'nombre').order_by('nombre')
            return JsonResponse([{'id': d['pk'], 'nombre': d['nombre']} for d in deptos], safe=False)
            
        elif action == 'buscar_carreras':
            depto_id = request.POST.get('id')
            carreras = Carrera.objects.filter(
                departamento_id=depto_id,
                departamento__facultad__universidad=universidad_obj
            ).values('pk', 'nombre').order_by('nombre')
            return JsonResponse([{'id': c['pk'], 'nombre': c['nombre']} for c in carreras], safe=False)
            
        # --- NUEVA API DE CHILE ---
        elif action == 'buscar_comunas':
            region_nombre = request.POST.get('id')
            chile_data = {
                "Arica y Parinacota": ["Arica", "Camarones", "Putre", "General Lagos"],
                "Tarapacá": ["Iquique", "Alto Hospicio", "Pozo Almonte", "Camiña", "Colchane", "Huara", "Pica"],
                "Antofagasta": ["Antofagasta", "Mejillones", "Sierra Gorda", "Taltal", "Calama", "Ollagüe", "San Pedro de Atacama", "Tocopilla", "María Elena"],
                "Atacama": ["Copiapó", "Caldera", "Tierra Amarilla", "Chañaral", "Diego de Almagro", "Vallenar", "Alto del Carmen", "Freirina", "Huasco"],
                "Coquimbo": ["La Serena", "Coquimbo", "Andacollo", "La Higuera", "Paiguano", "Vicuña", "Illapel", "Canela", "Los Vilos", "Salamanca", "Ovalle", "Combarbalá", "Monte Patria", "Punitaqui", "Río Hurtado"],
                "Valparaíso": ["Valparaíso", "Casablanca", "Concón", "Juan Fernández", "Puchuncaví", "Quintero", "Viña del Mar", "Isla de Pascua", "Los Andes", "Calle Larga", "Rinconada", "San Esteban", "La Ligua", "Cabildo", "Papudo", "Petorca", "Zapallar", "Quillota", "Calera", "Hijuelas", "La Cruz", "Nogales", "San Antonio", "Algarrobo", "Cartagena", "El Quisco", "El Tabo", "Santo Domingo", "San Felipe", "Catemu", "Llaillay", "Panquehue", "Putaendo", "Santa María", "Quilpué", "Limache", "Olmué", "Villa Alemana"],
                "Metropolitana": ["Cerrillos", "Cerro Navia", "Conchalí", "El Bosque", "Estación Central", "Huechuraba", "Independencia", "La Cisterna", "La Florida", "La Granja", "La Pintana", "La Reina", "Las Condes", "Lo Barnechea", "Lo Espejo", "Lo Prado", "Macul", "Maipú", "Ñuñoa", "Pedro Aguirre Cerda", "Peñalolén", "Providencia", "Pudahuel", "Quilicura", "Quinta Normal", "Recoleta", "Renca", "San Joaquín", "San Miguel", "San Ramón", "Santiago", "Vitacura", "Puente Alto", "Pirque", "San José de Maipo", "Colina", "Lampa", "Tiltil", "San Bernardo", "Buin", "Calera de Tango", "Paine", "Melipilla", "Alhué", "Curacaví", "María Pinto", "San Pedro", "Talagante", "El Monte", "Isla de Maipo", "Padre Hurtado", "Peñaflor"],
                "O'Higgins": ["Rancagua", "Codegua", "Coinco", "Coltauco", "Doñihue", "Graneros", "Las Cabras", "Machalí", "Malloa", "Mostazal", "Olivar", "Peumo", "Pichidegua", "Quinta de Tilcoco", "Rengo", "Requínoa", "San Vicente", "Pichilemu", "La Estrella", "Litueche", "Marchihue", "Navidad", "Paredones", "San Fernando", "Chépica", "Chimbarongo", "Lolol", "Nancagua", "Palmilla", "Peralillo", "Placilla", "Pumanque", "Santa Cruz"],
                "Maule": ["Talca", "Constitución", "Curepto", "Empedrado", "Maule", "Pelarco", "Pencahue", "Río Claro", "San Clemente", "San Rafael", "Cauquenes", "Chanco", "Pelluhue", "Curicó", "Hualañé", "Licantén", "Molina", "Rauco", "Romeral", "Sagrada Familia", "Teno", "Vichuquén", "Linares", "Colbún", "Longaví", "Parral", "Retiro", "San Javier", "Villa Alegre", "Yerbas Buenas"],
                "Ñuble": ["Chillán", "Bulnes", "Cobquecura", "Coelemu", "Coihueco", "Chillán Viejo", "El Carmen", "Ninhue", "Ñiquén", "Pemuco", "Pinto", "Portezuelo", "Quillón", "Quirihue", "Ránquil", "San Carlos", "San Fabián", "San Ignacio", "San Nicolás", "Treguaco", "Yungay"],
                "Biobío": ["Concepción", "Coronel", "Chiguayante", "Florida", "Hualqui", "Lota", "Penco", "San Pedro de la Paz", "Santa Juana", "Talcahuano", "Tomé", "Hualpén", "Lebu", "Arauco", "Cañete", "Contulmo", "Curanilahue", "Los Álamos", "Tirúa", "Los Ángeles", "Antuco", "Cabrero", "Laja", "Mulchén", "Nacimiento", "Negrete", "Quilaco", "Quilleco", "San Rosendo", "Santa Bárbara", "Tucapel", "Yumbel", "Alto Biobío"],
                "Araucanía": ["Temuco", "Carahue", "Cunco", "Curarrehue", "Freire", "Galvarino", "Gorbea", "Lautaro", "Loncoche", "Melipeuco", "Nueva Imperial", "Padre Las Casas", "Perquenco", "Pitrufquén", "Pucón", "Saavedra", "Teodoro Schmidt", "Toltén", "Vilcún", "Villarrica", "Cholchol", "Angol", "Collipulli", "Curacautín", "Ercilla", "Lonquimay", "Los Sauces", "Lumaco", "Purén", "Renaico", "Traiguén", "Victoria"],
                "Los Ríos": ["Valdivia", "Corral", "Lanco", "Los Lagos", "Máfil", "Mariquina", "Paillaco", "Panguipulli", "La Unión", "Futrono", "Lago Ranco", "Río Bueno"],
                "Los Lagos": ["Puerto Montt", "Calbuco", "Cochamó", "Fresia", "Frutillar", "Los Muermos", "Llanquihue", "Maullín", "Puerto Varas", "Castro", "Ancud", "Chonchi", "Curaco de Vélez", "Dalcahue", "Puqueldón", "Queilén", "Quellón", "Quemchi", "Quinchao", "Osorno", "Puerto Octay", "Purranque", "Puyehue", "Río Negro", "San Juan de la Costa", "San Pablo", "Chaitén", "Futaleufú", "Hualaihué", "Palena"],
                "Aysén": ["Coyhaique", "Lago Verde", "Aysén", "Cisnes", "Guaitecas", "Cochrane", "O'Higgins", "Tortel", "Chile Chico", "Río Ibáñez"],
                "Magallanes": ["Punta Arenas", "Laguna Blanca", "Río Verde", "San Gregorio", "Cabo de Hornos (Williams)", "Antártica", "Porvenir", "Primavera", "Timaukel", "Natales", "Torres del Paine"]
            }
            comunas = chile_data.get(region_nombre, [])
            return JsonResponse([{'id': c, 'nombre': c} for c in comunas], safe=False)

    # ---------------------------------------------------------
    # 3. BLOQUE POST (Crear Usuario)
    # ---------------------------------------------------------
    if request.method == "POST":
        # Datos de Cuenta
        username = request.POST.get("username", "").strip()
        email = request.POST.get("email", "").strip()
        
        # Datos Personales
        nombre = request.POST.get("nombre", "").strip()
        apellido = request.POST.get("apellido", "").strip()
        telefono = request.POST.get("telefono", "").strip()
        
        # --- NUEVOS CAMPOS ---
        region = request.POST.get("region")
        comuna = request.POST.get("comuna", "").strip()
        genero = request.POST.get("genero")
        fecha_nac = request.POST.get("fecha_nacimiento")
        # ---------------------
        
        # Datos Académicos
        carrera_id = request.POST.get("carrera")

        # Validaciones
        errores = []
        if not username: errores.append("El username es obligatorio.")
        if not email: errores.append("El email es obligatorio.")
        if not carrera_id: errores.append("Debes seleccionar una carrera.")
        
        if Usuario.objects.filter(username=username).exists():
            errores.append("El username ya existe.")

        if errores:
            for e in errores: messages.error(request, e)
            facultades = Facultad.objects.filter(universidad=universidad_obj)
            return render(request, "coordinador/crear_usuario.html", {
                "facultades": facultades, "base_template": "base_public.html", "volver_url": "coordina_usuarios"
            })

        try:
            # Generar Clave
            clave_temporal = get_random_string(length=8)
            
            # Crear Usuario
            nuevo_usuario = Usuario.objects.create_user(
                username=username,
                email=email,
                password=clave_temporal,
                first_name=nombre,
                last_name=apellido
            )
            
            # Guardar campos extra
            if telefono: nuevo_usuario.telefono = telefono
            if region: nuevo_usuario.region = region
            if comuna: nuevo_usuario.comuna = comuna
            if genero: nuevo_usuario.genero = genero
            if fecha_nac: nuevo_usuario.fecha_nacimiento = fecha_nac
            
            # Asignar Carrera
            carrera_obj = Carrera.objects.get(pk=carrera_id)
            nuevo_usuario.carrera = carrera_obj
            nuevo_usuario.save()
            
            # Crear Registro Histórico
            RegistroActividad.objects.create(
                usuario=nuevo_usuario,
                carrera=carrera_obj,
                descripcion="Creado por Coordinador"
            )

            # Enviar Correo
            uid = urlsafe_base64_encode(force_bytes(nuevo_usuario.pk))
            token = default_token_generator.make_token(nuevo_usuario)
            domain = request.get_host()
            scheme = request.scheme
            reset_url = reverse('password_reset_confirm', kwargs={'uidb64': uid, 'token': token})
            full_url = f"{scheme}://{domain}{reset_url}"
            asunto = "Bienvenido a Alumni - Activa tu cuenta"
            mensaje = f"Hola {nombre},\n\nTu cuenta ha sido creada en {universidad_obj.nombre}.\n\nUsuario: {username}\n\nPara activar tu cuenta y crear tu contraseña, haz clic aquí:\n{full_url}\n\nEste enlace expirará en 24 horas.\n\nBienvenido/a,\nEquipo Alumni"
            
            EmailThread(asunto, mensaje, [email]).start()
            
            messages.success(request, f"Alumno creado exitosamente. Correo de activación enviado a {email}.")
            return redirect("coordina_usuarios")

        except Exception as e:
            print(f"🔴 ERROR: {e}")
            messages.error(request, f"Error interno: {str(e)}")
    # ---------------------------------------------------------
    # 4. BLOQUE GET (Carga Inicial)
    # ---------------------------------------------------------
    facultades = Facultad.objects.filter(universidad=universidad_obj)

    return render(request, "coordinador/crear_usuario.html", {
        "facultades": facultades,
        "base_template": "base_public.html",
        "volver_url": "coordina_usuarios"
    })


@login_required
def ver_cv_coordinacion(request, user_id):
    coordinador = request.user

    # Verificar que tenga Universidad asignada
    if not coordinador.universidad_coordinador:
        return HttpResponseForbidden("No tienes una Universidad asignada.")

    # Obtener usuario a visualizar
    usuario = get_object_or_404(Usuario, id=user_id)

    # Verificar que pertenece a la Universidad del coordinador
    # Usamos la relación directa por carrera
    pertenece = False
    if usuario.carrera:
        if usuario.carrera.departamento.facultad.universidad == coordinador.universidad_coordinador:
            pertenece = True

    if not pertenece:
        return HttpResponseForbidden("No puedes ver el CV de un usuario de otra Universidad.")

    # Obtener CV
    cv = usuario.cvusuario_set.first()

    # Obtener palabras clave
    palabras = []
    if cv:
        palabras = [
            cv.palabra1, cv.palabra2, cv.palabra3, cv.palabra4, cv.palabra5,
            cv.palabra6, cv.palabra7, cv.palabra8, cv.palabra9, cv.palabra10,
        ]
        palabras = [p for p in palabras if p]

    # Renderizamos el template de Admin pero con diseño público
    return render(request, "coordinador/ver_cv.html", {
        "usuario": usuario,
        "cv": cv,
        "palabras": palabras,
        "base_template": "base_public.html",
        "volver_url": "coordina_usuarios"
    })

@login_required
def editar_usuario_coordinacion(request, user_id):
    usuario_a_editar = get_object_or_404(Usuario, id=user_id)
    universidad_coordinador = request.user.universidad_coordinador

    # --- API DE COMUNAS AJAX PARA EDICIÓN ---
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        action = request.POST.get('action')
        if action == 'buscar_comunas':
            region_nombre = request.POST.get('id')
            chile_data = {
                "Arica y Parinacota": ["Arica", "Camarones", "Putre", "General Lagos"],
                "Tarapacá": ["Iquique", "Alto Hospicio", "Pozo Almonte", "Camiña", "Colchane", "Huara", "Pica"],
                "Antofagasta": ["Antofagasta", "Mejillones", "Sierra Gorda", "Taltal", "Calama", "Ollagüe", "San Pedro de Atacama", "Tocopilla", "María Elena"],
                "Atacama": ["Copiapó", "Caldera", "Tierra Amarilla", "Chañaral", "Diego de Almagro", "Vallenar", "Alto del Carmen", "Freirina", "Huasco"],
                "Coquimbo": ["La Serena", "Coquimbo", "Andacollo", "La Higuera", "Paiguano", "Vicuña", "Illapel", "Canela", "Los Vilos", "Salamanca", "Ovalle", "Combarbalá", "Monte Patria", "Punitaqui", "Río Hurtado"],
                "Valparaíso": ["Valparaíso", "Casablanca", "Concón", "Juan Fernández", "Puchuncaví", "Quintero", "Viña del Mar", "Isla de Pascua", "Los Andes", "Calle Larga", "Rinconada", "San Esteban", "La Ligua", "Cabildo", "Papudo", "Petorca", "Zapallar", "Quillota", "Calera", "Hijuelas", "La Cruz", "Nogales", "San Antonio", "Algarrobo", "Cartagena", "El Quisco", "El Tabo", "Santo Domingo", "San Felipe", "Catemu", "Llaillay", "Panquehue", "Putaendo", "Santa María", "Quilpué", "Limache", "Olmué", "Villa Alemana"],
                "Metropolitana": ["Cerrillos", "Cerro Navia", "Conchalí", "El Bosque", "Estación Central", "Huechuraba", "Independencia", "La Cisterna", "La Florida", "La Granja", "La Pintana", "La Reina", "Las Condes", "Lo Barnechea", "Lo Espejo", "Lo Prado", "Macul", "Maipú", "Ñuñoa", "Pedro Aguirre Cerda", "Peñalolén", "Providencia", "Pudahuel", "Quilicura", "Quinta Normal", "Recoleta", "Renca", "San Joaquín", "San Miguel", "San Ramón", "Santiago", "Vitacura", "Puente Alto", "Pirque", "San José de Maipo", "Colina", "Lampa", "Tiltil", "San Bernardo", "Buin", "Calera de Tango", "Paine", "Melipilla", "Alhué", "Curacaví", "María Pinto", "San Pedro", "Talagante", "El Monte", "Isla de Maipo", "Padre Hurtado", "Peñaflor"],
                "O'Higgins": ["Rancagua", "Codegua", "Coinco", "Coltauco", "Doñihue", "Graneros", "Las Cabras", "Machalí", "Malloa", "Mostazal", "Olivar", "Peumo", "Pichidegua", "Quinta de Tilcoco", "Rengo", "Requínoa", "San Vicente", "Pichilemu", "La Estrella", "Litueche", "Marchihue", "Navidad", "Paredones", "San Fernando", "Chépica", "Chimbarongo", "Lolol", "Nancagua", "Palmilla", "Peralillo", "Placilla", "Pumanque", "Santa Cruz"],
                "Maule": ["Talca", "Constitución", "Curepto", "Empedrado", "Maule", "Pelarco", "Pencahue", "Río Claro", "San Clemente", "San Rafael", "Cauquenes", "Chanco", "Pelluhue", "Curicó", "Hualañé", "Licantén", "Molina", "Rauco", "Romeral", "Sagrada Familia", "Teno", "Vichuquén", "Linares", "Colbún", "Longaví", "Parral", "Retiro", "San Javier", "Villa Alegre", "Yerbas Buenas"],
                "Ñuble": ["Chillán", "Bulnes", "Cobquecura", "Coelemu", "Coihueco", "Chillán Viejo", "El Carmen", "Ninhue", "Ñiquén", "Pemuco", "Pinto", "Portezuelo", "Quillón", "Quirihue", "Ránquil", "San Carlos", "San Fabián", "San Ignacio", "San Nicolás", "Treguaco", "Yungay"],
                "Biobío": ["Concepción", "Coronel", "Chiguayante", "Florida", "Hualqui", "Lota", "Penco", "San Pedro de la Paz", "Santa Juana", "Talcahuano", "Tomé", "Hualpén", "Lebu", "Arauco", "Cañete", "Contulmo", "Curanilahue", "Los Álamos", "Tirúa", "Los Ángeles", "Antuco", "Cabrero", "Laja", "Mulchén", "Nacimiento", "Negrete", "Quilaco", "Quilleco", "San Rosendo", "Santa Bárbara", "Tucapel", "Yumbel", "Alto Biobío"],
                "Araucanía": ["Temuco", "Carahue", "Cunco", "Curarrehue", "Freire", "Galvarino", "Gorbea", "Lautaro", "Loncoche", "Melipeuco", "Nueva Imperial", "Padre Las Casas", "Perquenco", "Pitrufquén", "Pucón", "Saavedra", "Teodoro Schmidt", "Toltén", "Vilcún", "Villarrica", "Cholchol", "Angol", "Collipulli", "Curacautín", "Ercilla", "Lonquimay", "Los Sauces", "Lumaco", "Purén", "Renaico", "Traiguén", "Victoria"],
                "Los Ríos": ["Valdivia", "Corral", "Lanco", "Los Lagos", "Máfil", "Mariquina", "Paillaco", "Panguipulli", "La Unión", "Futrono", "Lago Ranco", "Río Bueno"],
                "Los Lagos": ["Puerto Montt", "Calbuco", "Cochamó", "Fresia", "Frutillar", "Los Muermos", "Llanquihue", "Maullín", "Puerto Varas", "Castro", "Ancud", "Chonchi", "Curaco de Vélez", "Dalcahue", "Puqueldón", "Queilén", "Quellón", "Quemchi", "Quinchao", "Osorno", "Puerto Octay", "Purranque", "Puyehue", "Río Negro", "San Juan de la Costa", "San Pablo", "Chaitén", "Futaleufú", "Hualaihué", "Palena"],
                "Aysén": ["Coyhaique", "Lago Verde", "Aysén", "Cisnes", "Guaitecas", "Cochrane", "O'Higgins", "Tortel", "Chile Chico", "Río Ibáñez"],
                "Magallanes": ["Punta Arenas", "Laguna Blanca", "Río Verde", "San Gregorio", "Cabo de Hornos (Williams)", "Antártica", "Porvenir", "Primavera", "Timaukel", "Natales", "Torres del Paine"]
            }
            comunas = chile_data.get(region_nombre, [])
            return JsonResponse([{'id': c, 'nombre': c} for c in comunas], safe=False)

    # Validar permisos: ¿El usuario pertenece a mi universidad?
    es_mi_alumno = False
    if usuario_a_editar.carrera:
        if usuario_a_editar.carrera.departamento.facultad.universidad == universidad_coordinador:
            es_mi_alumno = True
            
    if not es_mi_alumno:
        messages.error(request, "Este usuario no pertenece a tu universidad.")
        return redirect("coordina_usuarios")

    areas = Carrera.objects.filter(departamento__facultad__universidad=universidad_coordinador)

    if request.method == "POST":
        nombre = request.POST.get("nombre", "").strip()
        apellido = request.POST.get("apellido", "").strip()
        email = request.POST.get("email", "").strip()
        telefono = request.POST.get("telefono", "").strip()
        
        # --- CAMPOS NUEVOS ---
        region = request.POST.get("region")
        comuna = request.POST.get("comuna", "").strip()
        genero = request.POST.get("genero")
        fecha_nac = request.POST.get("fecha_nacimiento")
        # ---------------------
        
        area_id = request.POST.get("area")

        usuario_a_editar.first_name = nombre
        usuario_a_editar.last_name = apellido
        usuario_a_editar.email = email
        if hasattr(usuario_a_editar, 'telefono'):
            usuario_a_editar.telefono = telefono if telefono else None
            
        usuario_a_editar.region = region
        usuario_a_editar.comuna = comuna
        usuario_a_editar.genero = genero
        if fecha_nac: usuario_a_editar.fecha_nacimiento = fecha_nac
        
        if area_id:
            carrera_obj = Carrera.objects.get(pk=area_id)
            usuario_a_editar.carrera = carrera_obj
            
        usuario_a_editar.save()
        messages.success(request, "Usuario actualizado.")
        return redirect("coordina_usuarios")

    return render(request, "coordinador/editar_usuario.html", {
        "usuario": usuario_a_editar,
        "areas": areas,
        "base_template": "base_public.html",
        "volver_url": "coordina_usuarios"
    })


@login_required
def eliminar_usuario_coordinacion(request, user_id):
    usuario_a_borrar = get_object_or_404(Usuario, id=user_id)
    mi_universidad = request.user.universidad_coordinador

    if not mi_universidad:
        return redirect('home')

    # Validación de Pertenencia (Jerarquía)
    pertenece = False
    if usuario_a_borrar.carrera:
        if usuario_a_borrar.carrera.departamento.facultad.universidad == mi_universidad:
            pertenece = True
    
    if not pertenece:
        messages.error(request, "No puedes eliminar usuarios de otra universidad.")
        return redirect("coordina_usuarios")

    # Eliminar
    usuario_a_borrar.delete()
    messages.success(request, "Usuario eliminado correctamente.")
    return redirect("coordina_usuarios")

@login_required
def resetear_clave_coordinacion(request, user_id):
    if request.method != "POST" or not request.user.es_coordinador:
        return redirect("coordina_usuarios")

    usuario = get_object_or_404(Usuario, pk=user_id)
    mi_universidad = request.user.universidad_coordinador

    # Validación
    pertenece = False
    if usuario.carrera and usuario.carrera.departamento.facultad.universidad == mi_universidad:
        pertenece = True
        
    if not pertenece:
        messages.error(request, "No autorizado.")
        return redirect("coordina_usuarios")

    # Generar nueva clave
    nueva_clave = ''.join(random.choices(string.ascii_letters + string.digits, k=8))
    usuario.set_password(nueva_clave)
    usuario.save()

    try:
        send_mail(
            "Nueva Contraseña",
            f"Tu clave ha sido reseteada por coordinación.\nNueva clave: {nueva_clave}",
            settings.EMAIL_HOST_USER,
            [usuario.email],
            fail_silently=True
        )
        messages.success(request, f"Clave enviada a {usuario.email}")
    except:
        messages.warning(request, f"Clave cambiada a: {nueva_clave} (Correo falló)")

    return redirect("coordina_usuarios")


# --- ASEGÚRATE DE TENER ESTOS IMPORTS AL PRINCIPIO DE TU ARCHIVO views.py ---
from collections import Counter
from django.db.models import Q
from django.contrib import messages
from django.shortcuts import render, redirect
from django.contrib.auth.hashers import make_password
from django.core.mail import send_mail
from openpyxl import load_workbook
import random
import string

# Helper para generar clave si no lo tienes
def generar_clave(longitud=8):
    caracteres = string.ascii_letters + string.digits
    return ''.join(random.choice(caracteres) for _ in range(longitud))

# =================================================
# 1. PROCESAMIENTO DE CVS (Corregido y Blindado)
# =================================================
def procesamiento_cvs(request):
    try:
        Universidads = Universidad.objects.all()
        Universidad_id = request.GET.get("Universidad") or request.POST.get("Universidad")
        
        Universidad_seleccionada = None
        usuarios = []
        matriz_con_usuarios = []
        frecuencias = {}

        if Universidad_id:
            try:
                Universidad_seleccionada = Universidad.objects.get(id_universidad=Universidad_id)
            except Universidad.DoesNotExist:
                pass # Se queda en None

            if Universidad_seleccionada:
                # Búsqueda
                usuarios = Usuario.objects.filter(
                    Q(carrera__departamento__facultad__universidad=Universidad_seleccionada) |
                    Q(registroactividad__carrera__departamento__facultad__universidad=Universidad_seleccionada),
                    cvusuario__isnull=False
                ).distinct()

                # Si es POST, procesamos
                if request.method == "POST":
                    try:
                        lista_palabras = []
                        todas_las_palabras = []

                        # Extraer palabras
                        for usr in usuarios:
                            cv = CVUsuario.objects.filter(usuario=usr).first()
                            palabras = []
                            if cv:
                                raw_palabras = [
                                    cv.palabra1, cv.palabra2, cv.palabra3, cv.palabra4, cv.palabra5,
                                    cv.palabra6, cv.palabra7, cv.palabra8, cv.palabra9, cv.palabra10
                                ]
                                palabras = [str(p).strip().lower() for p in raw_palabras if p and str(p).strip()]
                            
                            lista_palabras.append(palabras)
                            todas_las_palabras.extend(palabras)

                        # Generar Matriz
                        tamaño = len(usuarios)
                        if tamaño > 0:
                            matriz = [[0] * tamaño for _ in range(tamaño)]
                            for i in range(tamaño):
                                for j in range(tamaño):
                                    if i == j:
                                        matriz[i][j] = 10 # Diagonal
                                        continue
                                    # Intersección
                                    coincidencias = len(set(lista_palabras[i]) & set(lista_palabras[j]))
                                    matriz[i][j] = coincidencias
                            
                            matriz_con_usuarios = list(zip(usuarios, matriz))

                        # Generar Frecuencias
                        if todas_las_palabras:
                            frecuencias = dict(Counter(todas_las_palabras).most_common(20))
                            
                    except Exception as e:
                        print(f"Error procesando datos CV: {e}")
                        messages.error(request, f"Error interno al calcular datos: {e}")

        return render(request, "procesamiento_cvs/procesamiento_cvs.html", {
            "Universidads": Universidads,
            "Universidad_seleccionada": Universidad_seleccionada,
            "usuarios": usuarios,
            "matriz_con_usuarios": matriz_con_usuarios,
            "frecuencias": frecuencias,
        })
    except Exception as e:
        print(f"Error CRÍTICO en vista CVs: {e}")
        messages.error(request, f"Error crítico: {e}")
        return redirect('home_admin')


# =================================================
# 2. PROCESAMIENTO OFREZCO/NECESITO (Corregido)
# =================================================
def procesamiento_ofrezco_necesito(request):
    try:
        Universidads = Universidad.objects.all()
        Universidad_id = request.GET.get("Universidad") or request.POST.get("Universidad")
        
        Universidad_seleccionada = None
        usuarios = []
        matriz_ofrezco_necesito = []
        frecuencias_ofrezco = {}
        frecuencias_necesito = {}

        if Universidad_id:
            try:
                Universidad_seleccionada = Universidad.objects.get(id_universidad=Universidad_id)
            except Universidad.DoesNotExist:
                pass

            if Universidad_seleccionada:
                usuarios = Usuario.objects.filter(
                    Q(carrera__departamento__facultad__universidad=Universidad_seleccionada) |
                    Q(registroactividad__carrera__departamento__facultad__universidad=Universidad_seleccionada)
                ).distinct()

                if request.method == "POST":
                    try:
                        lista_ofrezco = []
                        lista_necesito = []
                        todas_palabras_ofrezco = []
                        todas_palabras_necesito = []

                        for usr in usuarios:
                            # --- OFERTA ---
                            oferta = Oferta.objects.filter(usuario=usr).first()
                            ofrezco = []
                            if oferta:
                                campos = ["palabra1", "palabra2", "palabra3"]
                                if hasattr(oferta, 'palabra4'): campos.append('palabra4')
                                if hasattr(oferta, 'palabra5'): campos.append('palabra5')
                                for campo in campos:
                                    val = getattr(oferta, campo, None)
                                    if val and str(val).strip():
                                        val_clean = str(val).strip().lower()
                                        ofrezco.append(val_clean)
                                        todas_palabras_ofrezco.append(val_clean)
                            lista_ofrezco.append(ofrezco)

                            # --- NECESIDAD ---
                            necesidad = Necesidad.objects.filter(usuario=usr).first()
                            necesito = []
                            if necesidad:
                                campos = ["palabra1", "palabra2", "palabra3"]
                                if hasattr(necesidad, 'palabra4'): campos.append('palabra4')
                                if hasattr(necesidad, 'palabra5'): campos.append('palabra5')
                                for campo in campos:
                                    val = getattr(necesidad, campo, None)
                                    if val and str(val).strip():
                                        val_clean = str(val).strip().lower()
                                        necesito.append(val_clean)
                                        todas_palabras_necesito.append(val_clean)
                            lista_necesito.append(necesito)

                        # --- MATRIZ ---
                        n = len(usuarios)
                        if n > 0:
                            matriz = [[0] * n for _ in range(n)]
                            for i in range(n):
                                for j in range(n):
                                    # Cruce: Lo que 'i' ofrece vs lo que 'j' necesita
                                    coincidencias = len(set(lista_ofrezco[i]) & set(lista_necesito[j]))
                                    matriz[i][j] = coincidencias
                            matriz_ofrezco_necesito = list(zip(usuarios, matriz))

                        # --- FRECUENCIAS ---
                        if todas_palabras_ofrezco:
                            frecuencias_ofrezco = dict(Counter(todas_palabras_ofrezco).most_common(20))
                        if todas_palabras_necesito:
                            frecuencias_necesito = dict(Counter(todas_palabras_necesito).most_common(20))
                            
                    except Exception as e:
                        print(f"Error procesando matriz O/N: {e}")
                        messages.error(request, f"Error al procesar datos: {e}")

        return render(request, "procesamiento_ofrezco_necesito/procesamiento_ofrezco_necesito.html", {
            "Universidads": Universidads,
            "Universidad_seleccionada": Universidad_seleccionada,
            "usuarios": usuarios,
            "matriz_ofrezco_necesito": matriz_ofrezco_necesito,
            "frecuencias_ofrezco": frecuencias_ofrezco,
            "frecuencias_necesito": frecuencias_necesito,
        })
    except Exception as e:
        print(f"Error CRÍTICO en vista O/N: {e}")
        messages.error(request, f"Error crítico: {e}")
        return redirect('home_admin')


# =================================================
# 3. CARGAR EXCEL USUARIOS (Corregido)
# =================================================
@login_required
def cargar_excel_usuarios(request):
    try:
        usuario_logueado = request.user
        
        # 1. Determinar Rol y Contexto Visual
        es_coordinador = getattr(usuario_logueado, 'es_coordinador', False) and usuario_logueado.universidad_coordinador is not None
        
        # Variables por defecto para el template
        base_template = 'base_admin.html'
        volver_url = 'admin_usuarios'
        
        Universidads = []
        Universidad_seleccionada = None

        if es_coordinador:
            # COORDINADOR: Su universidad es fija
            base_template = 'base_public.html' # O el que uses para coord
            volver_url = 'coordina_usuarios'
            Universidad_seleccionada = usuario_logueado.universidad_coordinador
            Universidads = [Universidad_seleccionada] # Lista de 1 elemento
        else:
            # ADMIN: Puede elegir universidad
            Universidads = Universidad.objects.all()
            # Obtener ID por GET (al entrar) o mantenerlo si ya estaba seleccionado
            uni_id = request.GET.get("universidad") or request.GET.get("Universidad")
            
            if uni_id:
                try:
                    Universidad_seleccionada = Universidad.objects.get(pk=uni_id)
                except Universidad.DoesNotExist:
                    pass

        # Contexto base para renderizar (se usa tanto en GET como en errores de POST)
        context = {
            "Universidads": Universidads, 
            "Universidad_seleccionada": Universidad_seleccionada,
            "es_coordinador_view": es_coordinador,
            "base_template": base_template,
            "volver_url": volver_url
        }

        # 2. Validación Inicial
        if not Universidad_seleccionada:
            # Si es admin y no ha elegido U, mostramos el selector del template
            return render(request, "admin_usuarios/cargar_excel.html", context)

        # 3. Procesamiento del Archivo (POST)
        if request.method == "POST":
            archivo = request.FILES.get("archivo")
            if not archivo:
                messages.error(request, "Debe seleccionar un archivo Excel.")
                return render(request, "admin_usuarios/cargar_excel.html", context)

            try:
                wb = load_workbook(archivo)
                ws = wb.active
            except Exception as e:
                messages.error(request, f"Error al leer el archivo: {e}")
                return render(request, "admin_usuarios/cargar_excel.html", context)

            # Leer encabezados
            filas = list(ws.iter_rows(values_only=True))
            if not filas:
                messages.error(request, "El archivo está vacío.")
                return render(request, "admin_usuarios/cargar_excel.html", context)
            
            # Normalizar headers a minúsculas para evitar errores por "Email" vs "email"
            columnas = [str(c).lower().strip() for c in filas[0] if c] 
            
            # Mapeo de columnas requeridas
            req = ["username", "first_name", "last_name", "email", "area"]
            faltantes = [r for r in req if r not in columnas]
            
            if faltantes:
                messages.error(request, f"Faltan columnas en el Excel: {', '.join(faltantes)}")
                return render(request, "admin_usuarios/cargar_excel.html", context)

            # Procesar datos
            creados = 0
            duplicados = 0
            
            # Mapa de índices (dinámico, no importa el orden de columnas)
            idx = {name: columnas.index(name) for name in req}
            idx_tel = columnas.index("telefono") if "telefono" in columnas else -1

            for row in filas[1:]:
                if not row: continue # Saltar filas vacías
                
                try:
                    # Extracción segura de datos
                    def get_val(i): return str(row[i]).strip() if i < len(row) and row[i] else ""
                    
                    username = get_val(idx["username"])
                    first = get_val(idx["first_name"])
                    last = get_val(idx["last_name"])
                    email = get_val(idx["email"])
                    area_nombre = get_val(idx["area"])
                    telefono = get_val(idx_tel) if idx_tel != -1 else None

                    # Validación básica de fila
                    if not username or not email or not area_nombre: 
                        duplicados += 1
                        continue

                    # Validación de duplicados en BD
                    if Usuario.objects.filter(Q(username=username) | Q(email=email)).exists():
                        duplicados += 1 # Ya existe
                        continue

                    # Buscar carrera/área en la Universidad Seleccionada
                    area = Carrera.objects.filter(
                        nombre__iexact=area_nombre,
                        departamento__facultad__universidad=Universidad_seleccionada
                    ).first()

                    if not area:
                        # Si la carrera no existe en ESA universidad, saltamos
                        print(f"⚠️ Área '{area_nombre}' no encontrada en {Universidad_seleccionada.nombre}")
                        duplicados += 1 
                        continue

                    # Crear Usuario
                    clave = generar_clave()
                    user = Usuario.objects.create(
                        username=username,
                        first_name=first,
                        last_name=last,
                        email=email,
                        telefono=telefono,
                        carrera=area,
                        password=make_password(clave),
                        is_active=True
                    )
                    
                    # Vincular con Carrera (RegistroActividad)
                    RegistroActividad.objects.create(usuario=user, carrera=area)

                    # Enviar correo en SEGUNDO PLANO (Usando tu clase EmailThread)
                    uid = urlsafe_base64_encode(force_bytes(user.pk))
                    token = default_token_generator.make_token(user)
                    domain = request.get_host()
                    scheme = request.scheme
                    reset_url = reverse('password_reset_confirm', kwargs={'uidb64': uid, 'token': token})
                    full_url = f"{scheme}://{domain}{reset_url}"
                    asunto = "Bienvenido a Alumni - Activa tu cuenta"
                    mensaje = f"Hola {first},\n\nTu cuenta ha sido creada exitosamente en {Universidad_seleccionada.nombre}.\n\nUsuario: {username}\n\nPara activar tu cuenta y crear tu contraseña, haz clic aquí:\n{full_url}\n\nEste enlace expirará en 24 horas.\n\nBienvenido/a,\nEquipo Alumni"
                    
                    EmailThread(asunto, mensaje, [email]).start()

                    creados += 1

                except Exception as e:
                    print(f"❌ Error procesando fila Excel: {e}")
                    duplicados += 1

            messages.success(request, f"Carga completada. Usuarios creados: {creados}. Ignorados/Errores: {duplicados}.")

    except Exception as e:
        messages.error(request, f"Error crítico en el servidor: {e}")

    return render(request, "admin_usuarios/cargar_excel.html", context)

# Helper para contexto rápido
def ctx_excel(unis, sel):
    return {"Universidads": unis, "Universidad_seleccionada": sel}

class EmailThread(threading.Thread):
    def __init__(self, subject, message, recipient_list):
        self.subject = subject
        self.message = message
        self.recipient_list = recipient_list
        threading.Thread.__init__(self)

    def run(self):
        try:
            import smtplib
            import ssl
            from email.mime.text import MIMEText
            
            context = ssl.create_default_context()
            context.check_hostname = False
            context.verify_mode = ssl.CERT_NONE
            
            msg = MIMEText(self.message)
            msg['Subject'] = self.subject
            msg['From'] = settings.DEFAULT_FROM_EMAIL
            msg['To'] = ', '.join(self.recipient_list)
            
            with smtplib.SMTP('smtp-relay.brevo.com', 587) as server:
                server.ehlo()
                server.starttls(context=context)
                server.login('a4ba9b001@smtp-brevo.com', 'UPLM92nEtK50FTcD')
                server.sendmail(settings.DEFAULT_FROM_EMAIL, self.recipient_list, msg.as_string())
            
            print(f"✅ [BREVO] Correo enviado a {self.recipient_list}")
        except Exception as e:
            print(f"❌ [ERROR BREVO] {e}")



def link_callback(uri, rel):
    """
    Ayuda a xhtml2pdf a encontrar imágenes y CSS estáticos.
    """
    # Intentar encontrar con finders de Django
    result = finders.find(uri)
    if result:
        if not isinstance(result, (list, tuple)):
            result = [result]
        result = list(os.path.realpath(path) for path in result)
        path = result[0]
    else:
        # Si falla, intentar construir ruta manualmente
        sUrl = settings.STATIC_URL        
        sRoot = settings.STATIC_ROOT      
        mUrl = settings.MEDIA_URL         
        mRoot = settings.MEDIA_ROOT       

        if uri.startswith(mUrl):
            path = os.path.join(mRoot, uri.replace(mUrl, ""))
        elif uri.startswith(sUrl):
            path = os.path.join(sRoot, uri.replace(sUrl, ""))
        else:
            return uri

    # Verificar existencia
    if not os.path.isfile(path):
        print(f"⚠️ xhtml2pdf no encuentra: {path}")
        return uri 
        
    return path


@login_required
def gestionar_asistencia(request, evento_id):
    evento = get_object_or_404(Evento, pk=evento_id)
    usuario_logueado = request.user
    
    # Validación de permisos
    permitido = False
    base_template = 'base_admin.html'
    volver_url = 'home_admin' 

    if usuario_logueado.is_staff:
        permitido = True
    elif hasattr(usuario_logueado, 'universidad_coordinador') and usuario_logueado.universidad_coordinador == evento.universidad:
        permitido = True
        base_template = 'base_public.html'
        volver_url = 'panel_coordinador'
        
    if not permitido:
        messages.error(request, "No tienes permisos para gestionar este evento.")
        return redirect('home_admin')

    # Auto-poblado de invitados
    filtros_base = Q(is_active=True, is_staff=False)
    alumnos_objetivo = Usuario.objects.none()

    if evento.carrera:
        alumnos_objetivo = Usuario.objects.filter(filtros_base, carrera=evento.carrera)
    elif evento.departamento:
        alumnos_objetivo = Usuario.objects.filter(filtros_base, carrera__departamento=evento.departamento)
    elif evento.facultad:
        alumnos_objetivo = Usuario.objects.filter(filtros_base, carrera__departamento__facultad=evento.facultad)
    else:
        alumnos_objetivo = Usuario.objects.filter(filtros_base, carrera__departamento__facultad__universidad=evento.universidad)

    nuevas_creadas = 0
    for alumno in alumnos_objetivo:
        _, created = Invitacion.objects.get_or_create(
            evento=evento, usuario=alumno, defaults={'estado': 'SIN_ENVIAR'}
        )
        if created: nuevas_creadas += 1

    if nuevas_creadas > 0:
        messages.info(request, f"Se han añadido {nuevas_creadas} alumnos.")

    # Datos
    invitaciones = Invitacion.objects.filter(evento=evento).select_related('usuario', 'usuario__carrera')
    
    # Stats manuales
    stats = {
        'total': invitaciones.count(),
        'confirmados': invitaciones.filter(estado__in=['CONFIRMADO', 'ENTRADA_ENVIADA']).count(),
        'pendientes': invitaciones.filter(estado='ENVIADO').count(),
        'rechazados': invitaciones.filter(estado='RECHAZADO').count(),
        'por_enviar': invitaciones.filter(estado='SIN_ENVIAR').count(),
    }

    context = {
        'evento': evento,
        'invitaciones': invitaciones,
        'confirmados': invitaciones.filter(estado__in=['CONFIRMADO', 'ENTRADA_ENVIADA']),
        'stats': stats,
        'base_template': base_template,
        'volver_url': volver_url
    }
    return render(request, 'eventos/gestion_asistencia.html', context)


# 3. REPORTE CONFIRMADOS
@login_required
def reporte_confirmados(request, evento_id):
    evento = get_object_or_404(Evento, pk=evento_id)
    usuario_logueado = request.user
    
    # 1. Definir permisos y botón "Volver"
    base_template = 'base_admin.html'
    # Por defecto volvemos al home, pero lo ideal es modificar el template
    # para volver a 'gestion_asistencia' (ver nota abajo)
    volver_url = 'home_admin' 

    if usuario_logueado.is_staff:
        pass # Es admin
    elif hasattr(usuario_logueado, 'universidad_coordinador') and usuario_logueado.universidad_coordinador == evento.universidad:
        base_template = 'base_public.html'
        volver_url = 'panel_coordinador'
    else:
        messages.error(request, "No tienes permisos.")
        return redirect('home_admin')

    # 2. Obtener solo los confirmados
    confirmados = Invitacion.objects.filter(
        evento=evento,
        estado__in=['CONFIRMADO', 'ENTRADA_ENVIADA']
    ).select_related('usuario', 'usuario__carrera')

    context = {
        'evento': evento,
        'confirmados': confirmados,
        'base_template': base_template,
        'volver_url': volver_url
    }
    
    # Asegúrate que el archivo esté en templates/eventos/
    return render(request, 'eventos/reporte_confirmados.html', context)


# 4. DESCARGAR PDF
@login_required
def descargar_pdf_confirmados(request, evento_id):
    try:
        evento = get_object_or_404(Evento, pk=evento_id)
        
        # Validar permisos
        es_coordinador = hasattr(request.user, 'universidad_coordinador') and request.user.universidad_coordinador == evento.universidad
        if not request.user.is_staff and not es_coordinador:
            return redirect('home_admin')

        confirmados = Invitacion.objects.filter(
            evento=evento,
            estado__in=['CONFIRMADO', 'ENTRADA_ENVIADA']
        ).select_related('usuario', 'usuario__carrera')

        if not confirmados.exists():
            return HttpResponse("<script>alert('No hay confirmados para generar PDF');window.history.back();</script>")

        # Asegúrate que la ruta coincida con donde guardaste el HTML del PDF
        template_path = 'eventos/pdf_asistentes.html'
        
        context = {
            'confirmados': confirmados, 
            'evento': evento, 
            'STATIC_URL': settings.STATIC_URL
        }
        
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="asistentes_{evento.pk}.pdf"'
        
        template = get_template(template_path)
        html = template.render(context)

        pisa_status = pisa.CreatePDF(html, dest=response)

        if pisa_status.err:
            return HttpResponse(f'Error PDF: {pisa_status.err}')
        return response

    except Exception as e:
        import traceback
        return HttpResponse(f"Error: {e}")
    
# 5. ENVIAR RECORDATORIOS
@login_required
def enviar_recordatorio_pendientes(request, evento_id):
    evento = get_object_or_404(Evento, pk=evento_id)

    # Validar permisos
    es_admin = request.user.is_staff
    es_coordinador = hasattr(request.user, 'universidad_coordinador') and request.user.universidad_coordinador == evento.universidad
    
    if not (es_admin or es_coordinador):
        messages.error(request, "No tienes permisos.")
        return redirect('home_admin')

    pendientes = Invitacion.objects.filter(evento=evento, estado='ENVIADO')
    
    if not pendientes.exists():
        messages.warning(request, "No hay usuarios pendientes para recordar.")
        # OJO: Asegúrate que el nombre sea 'gestionar_asistencia' o 'gestion_asistencia' según tu urls.py
        return redirect('gestionar_asistencia', evento_id=evento.pk)

    mensajes_email = []
    # Usamos la URL que configuramos para las invitaciones
    link = request.build_absolute_uri('/eventos/mis-invitaciones/') 

    for invitacion in pendientes:
        if invitacion.usuario.email:
            asunto = f"RECORDATORIO: {evento.titulo}"
            cuerpo = f"""
                Hola {invitacion.usuario.first_name},
                
                Te recordamos confirmar tu asistencia para el evento "{evento.titulo}".
                
                📅 Fecha: {evento.inicio.strftime('%d/%m/%Y a las %H:%M')}
                
                Ingresa aquí para confirmar o rechazar: {link}
                
                Saludos, 
                Equipo {evento.universidad.nombre}
            """
            email = EmailMessage(
                asunto, 
                textwrap.dedent(cuerpo).strip(), 
                settings.DEFAULT_FROM_EMAIL, # Usamos el remitente configurado
                [invitacion.usuario.email]
            )
            mensajes_email.append(email)

    try:
        if mensajes_email:
            connection = get_connection()
            # Con Brevo/Anymail basta con enviar la lista de mensajes
            enviados = connection.send_messages(mensajes_email)
            messages.success(request, f"¡Éxito! Se enviaron {enviados} recordatorios.")
        else:
            messages.info(request, "No se encontraron correos válidos para enviar.")
    except Exception as e:
        print(f"Error Anymail: {e}") # Esto aparecerá en los logs de Render
        messages.error(request, f"Error al enviar correos: {str(e)}")
    
    # IMPORTANTE: Revisa que este nombre coincida con tu path de urls.py
    return redirect('gestionar_asistencia', evento_id=evento.pk)


@login_required
def calendario_admin_view(request):
    es_admin = request.user.is_staff
    es_coordinador = hasattr(request.user, 'universidad_coordinador') and request.user.universidad_coordinador
    
    if not (es_admin or es_coordinador):
        return redirect('ver_calendario') # Expulsa a los alumnos curiosos

    id_empresa = 0
    nombre = "Administración"
    uni_id_get = request.GET.get('empresa')

    if es_admin and uni_id_get:
        universidad_obj = get_object_or_404(Universidad, pk=uni_id_get)
        id_empresa = universidad_obj.pk
        nombre = universidad_obj.nombre
        
    elif es_coordinador:
        id_empresa = request.user.universidad_coordinador.id
        nombre = request.user.universidad_coordinador.nombre

    context = {
        'base_template': 'base_admin.html', 
        'es_editor': True, # <--- ESTO FUERZA MODO EDICIÓN
        'id_empresa': id_empresa,
        'nombre_empresa': nombre,
        'volver_url': 'home_admin' if es_admin else 'panel_coordinador'
    }
    
    # Renderiza la vista del gestor (Asegúrate de que la ruta de la carpeta sea correcta)
    return render(request, 'eventos/calendario_admin.html', context)




@login_required
def detalle_usuario_admin(request, user_id):
    # Verificamos que sea admin o coordinador
    if not request.user.is_staff and not request.user.es_coordinador:
        messages.error(request, "No tienes permisos para ver esto.")
        return redirect('home')

    usuario_det = get_object_or_404(Usuario, id=user_id)
    trabajos_empresas = TrabajoEmpresa.objects.filter(usuario=usuario_det).order_by('-id')
    
    # Obtener palabras clave de ofrezco/necesito
    ultima_oferta = Oferta.objects.filter(creado_por=usuario_det).last()
    palabras_oferta = [getattr(ultima_oferta, f'palabra{i}') for i in range(1, 6) if ultima_oferta and getattr(ultima_oferta, f'palabra{i}')]
    
    ultima_necesidad = Necesidad.objects.filter(usuario=usuario_det).last()
    palabras_necesidad = [getattr(ultima_necesidad, f'palabra{i}') for i in range(1, 6) if ultima_necesidad and getattr(ultima_necesidad, f'palabra{i}')]

    # Buscar usuarios similares basados en la IA de Gemini (Categorías iguales)
    categorias_usuario = trabajos_empresas.values_list('categoria_ia', flat=True).distinct()
    usuarios_similares = []
    if categorias_usuario:
        # Busca usuarios (distintos al actual) que tengan trabajos en la misma categoría
        usuarios_similares = Usuario.objects.filter(
            trabajos__categoria_ia__in=categorias_usuario
        ).exclude(id=usuario_det.id).distinct()[:6] # Mostramos máximo 6

    # Si es coordinador, mantendremos su layout
    base_template = 'base_admin.html' if request.user.is_staff else 'base_public.html'

    # CORRECCIÓN AQUÍ: Agregamos la subcarpeta 'admin_usuarios/' a la ruta
    return render(request, 'admin_usuarios/detalle_usuario_admin.html', {
        'usuario_det': usuario_det,
        'trabajos_empresas': trabajos_empresas,
        'palabras_oferta': palabras_oferta,
        'palabras_necesidad': palabras_necesidad,
        'usuarios_similares': usuarios_similares,
        'base_template': base_template
    })

def obtener_categoria_con_ia(nombre_empresa, descripcion):
    """
    Usa Gemini para analizar la empresa y asignar una categoría general.
    """
    try:
        if not hasattr(settings, 'GOOGLE_API_KEY') or not settings.GOOGLE_API_KEY:
            print("🚨 GEMINI ERROR: No se encontró GOOGLE_API_KEY en settings.")
            return "General"

        genai.configure(api_key=settings.GOOGLE_API_KEY) 
        
        # --- LA MAGIA DE TU OTRA FUNCIÓN (AUTO-DETECCIÓN) ---
        mis_modelos = []
        try:
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    mis_modelos.append(m.name)
        except Exception as e:
            print(f"⚠️ Error listando modelos: {e}")

        # Si no encuentra nada, cortamos por lo sano
        if not mis_modelos:
            return "General"

        # Priorizamos 'flash', si no, el primero que encuentre
        modelo_final = next((m for m in mis_modelos if 'flash' in m), mis_modelos[0])
        print(f"🚀 [Admin] Usando modelo para categoría: {modelo_final}")
        
        # Inicializamos con el modelo correcto que Google nos acaba de dar
        model = genai.GenerativeModel(modelo_final)
        # ----------------------------------------------------
        
        prompt = f"""
        Actúa como un clasificador de industrias para una red de networking. 
        Lee la siguiente empresa y descríbela en UNA sola categoría general (máximo 2 palabras).
        Ejemplos de salida: Tecnología, Salud, Gastronomía, Educación, Logística, Marketing.
        
        Empresa: {nombre_empresa}
        Descripción: {descripcion}
        
        Categoría:"""
        
        respuesta = model.generate_content(prompt)
        return respuesta.text.strip().title()
        
    except Exception as e:
        print("="*50)
        print(f"🚨 ERROR CRÍTICO CON GEMINI: {str(e)}")
        print("="*50)
        return "General"
    


@login_required
def admin_guardar_empresa_trabajo(request, user_id):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')
        
    usuario_afectado = get_object_or_404(Usuario, id=user_id)
    
    if request.method == 'POST':
        tipo = request.POST.get('tipo')
        nombre_empresa = request.POST.get('nombre_empresa')
        url = request.POST.get('url')
        email = request.POST.get('email')
        telefono = request.POST.get('telefono')
        descripcion_breve = request.POST.get('descripcion_breve')
        
        categoria_ia = obtener_categoria_con_ia(nombre_empresa, descripcion_breve)

        TrabajoEmpresa.objects.create(
            usuario=usuario_afectado,
            tipo=tipo,
            nombre_empresa=nombre_empresa,
            url=url,
            email=email,
            telefono=telefono,
            descripcion_breve=descripcion_breve,
            categoria_ia=categoria_ia
        )
        messages.success(request, f'Empresa agregada al usuario {usuario_afectado.username}.')
        
    return redirect('detalle_usuario_admin', user_id=user_id)

@login_required
def admin_eliminar_empresa_trabajo(request, id):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')
        
    trabajo = get_object_or_404(TrabajoEmpresa, id=id)
    user_id = trabajo.usuario.id
    
    if request.method == 'POST':
        trabajo.delete()
        messages.success(request, 'Registro eliminado correctamente.')
        
    return redirect('detalle_usuario_admin', user_id=user_id)

@login_required
def admin_preferencias(request):
    if not request.user.is_staff and not request.user.es_coordinador:
        messages.error(request, "No tienes permisos para ver esta página.")
        return redirect('home')

    from user.models import Preferencia
    from app.models import Universidad

    universidades = Universidad.objects.all()
    univ_id = request.GET.get('universidad')
    universidad_seleccionada = None

    if request.user.is_staff and not request.user.universidad_coordinador:
        # Es SuperAdmin: Ve el selector y filtramos
        if univ_id and univ_id.isdigit():
            universidad_seleccionada = int(univ_id)
        elif universidades.exists():
            # Selecciona la primera por defecto si no hay nada en la URL
            universidad_seleccionada = universidades.first().id_universidad
            
        preferencias = Preferencia.objects.filter(universidad_id=universidad_seleccionada).order_by('id')
    else:
        # Es Coordinador: Solo ve su universidad
        univ_coord = request.user.universidad_coordinador
        universidad_seleccionada = univ_coord.id_universidad if univ_coord else None
        preferencias = Preferencia.objects.filter(universidad=univ_coord).order_by('id')
        
        # Filtramos el selector para que el coordinador solo vea la suya
        if univ_coord:
            universidades = Universidad.objects.filter(id_universidad=univ_coord.id_universidad)

    base_template = 'base_admin.html' if request.user.is_staff else 'base_public.html'

    return render(request, 'admin_usuarios/admin_preferencias.html', {
        'preferencias': preferencias,
        'universidades': universidades,
        'universidad_seleccionada': universidad_seleccionada,
        'base_template': base_template
    })



@login_required
def admin_guardar_preferencia(request):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')

    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        pregunta = request.POST.get('pregunta')
        icono = request.POST.get('icono', 'fa-tag')
        univ_id_post = request.POST.get('universidad_id') # Leemos en qué pestaña estaba

        if nombre and pregunta:
            from user.models import Preferencia
            from app.models import Universidad
            
            # Determinamos a qué universidad se la asignamos
            univ_asignar = None
            if request.user.universidad_coordinador:
                univ_asignar = request.user.universidad_coordinador
            elif univ_id_post and univ_id_post.isdigit():
                univ_asignar = Universidad.objects.filter(id_universidad=univ_id_post).first()

            Preferencia.objects.create(
                nombre=nombre, 
                pregunta=pregunta, 
                icono=icono,
                universidad=univ_asignar
            )
            messages.success(request, f'Preferencia "{nombre}" agregada correctamente.')
        else:
            messages.error(request, 'Faltan datos obligatorios.')

    # Redirigimos a la misma pestaña en la que estaba
    url = reverse('admin_preferencias')
    if request.POST.get('universidad_id'):
        url += f"?universidad={request.POST.get('universidad_id')}"
    return redirect(url)



@login_required
def guardar_preferencias_usuario(request):
    if request.method == 'POST':
        from user.models import Preferencia
        
        # Capturamos todos los IDs de los checkboxes que el usuario marcó
        preferencias_ids = request.POST.getlist('preferencias_seleccionadas')
        
        # Limpiamos las preferencias anteriores
        request.user.preferencias.clear()
        
        if preferencias_ids:
            # SEGURIDAD CRÍTICA: Verificamos que las IDs enviadas pertenezcan a la universidad del alumno
            if request.user.carrera:
                mi_univ = request.user.carrera.departamento.facultad.universidad
                # Filtramos las IDs válidas
                preferencias_validas = Preferencia.objects.filter(id__in=preferencias_ids, universidad=mi_univ).values_list('id', flat=True)
                
                # Guardamos solo las que pasaron el filtro de seguridad
                request.user.preferencias.add(*preferencias_validas)

        messages.success(request, 'Tus preferencias han sido guardadas y actualizadas correctamente.')
        
    return redirect('mi_perfil')


@login_required
def admin_editar_preferencia(request, pref_id):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')

    if request.method == 'POST':
        from user.models import Preferencia 
        from django.shortcuts import get_object_or_404
        
        preferencia = get_object_or_404(Preferencia, id=pref_id)

        if not request.user.is_staff and preferencia.universidad != request.user.universidad_coordinador:
            messages.error(request, "No puedes editar preferencias de otra universidad.")
            return redirect('admin_preferencias')
        
        nuevo_nombre = request.POST.get('nombre')
        nueva_pregunta = request.POST.get('pregunta')
        nuevo_icono = request.POST.get('icono')

        if nuevo_nombre and nueva_pregunta:
            preferencia.nombre = nuevo_nombre
            preferencia.pregunta = nueva_pregunta
            if nuevo_icono:
                preferencia.icono = nuevo_icono
            
            preferencia.save()
            messages.success(request, f'Preferencia "{nuevo_nombre}" actualizada correctamente.')
        else:
            messages.error(request, 'El nombre y la pregunta son obligatorios.')

    # Redirigimos a la misma pestaña en la que estaba
    url = reverse('admin_preferencias')
    if request.POST.get('universidad_id'):
        url += f"?universidad={request.POST.get('universidad_id')}"
    return redirect(url)

@login_required
def admin_editar_empresa_trabajo(request, trabajo_id):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')
        
    from user.models import TrabajoEmpresa # Ajusta esto si tu modelo está en 'app.models'
    
    trabajo = get_object_or_404(TrabajoEmpresa, id=trabajo_id)
    usuario_id = trabajo.usuario.id
    
    if request.method == 'POST':
        nuevo_nombre = request.POST.get('nombre_empresa', trabajo.nombre_empresa)
        nueva_desc = request.POST.get('descripcion_breve', trabajo.descripcion_breve)
        
        # Si tienes importada la función de IA en views.py, puedes descomentar esto:
        # if nuevo_nombre != trabajo.nombre_empresa or nueva_desc != trabajo.descripcion_breve:
        #     trabajo.categoria_ia = obtener_categoria_con_ia(nuevo_nombre, nueva_desc)
            
        trabajo.tipo = request.POST.get('tipo', trabajo.tipo)
        trabajo.nombre_empresa = nuevo_nombre
        trabajo.descripcion_breve = nueva_desc
        trabajo.url = request.POST.get('url', trabajo.url)
        trabajo.email = request.POST.get('email', trabajo.email)
        trabajo.telefono = request.POST.get('telefono', trabajo.telefono)
        trabajo.whatsapp = request.POST.get('whatsapp', trabajo.whatsapp)
        
        trabajo.save()
        messages.success(request, f'Registro "{trabajo.nombre_empresa}" actualizado correctamente.')
        
    return redirect('detalle_usuario_admin', user_id=usuario_id)


@login_required
def admin_eliminar_preferencia(request, pref_id):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')

    if request.method == 'POST':
        from user.models import Preferencia 
        from django.shortcuts import get_object_or_404
        
        preferencia = get_object_or_404(Preferencia, id=pref_id)

        # Seguridad: Solo puede borrar si es su universidad o si es Superadmin
        if not request.user.is_staff and preferencia.universidad != request.user.universidad_coordinador:
            messages.error(request, "No puedes eliminar preferencias de otra universidad.")
        else:
            nombre = preferencia.nombre
            preferencia.delete()
            messages.success(request, f'Preferencia "{nombre}" eliminada correctamente.')

    # Redirigimos a la misma pestaña en la que estaba
    url = reverse('admin_preferencias')
    if request.POST.get('universidad_id'):
        url += f"?universidad={request.POST.get('universidad_id')}"
    return redirect(url)





# --- VISTAS PARA DESCUENTOS (Bancos) ---

@login_required
def admin_descuentos(request):
    if not request.user.is_staff and not request.user.es_coordinador:
        messages.error(request, "No tienes permisos.")
        return redirect('home')

    from user.models import DescuentoBanco
    from app.models import Universidad

    universidades = Universidad.objects.all()
    univ_id = request.GET.get('universidad')
    universidad_seleccionada = None

    if request.user.is_staff and not request.user.universidad_coordinador:
        if univ_id and univ_id.isdigit():
            universidad_seleccionada = int(univ_id)
        elif universidades.exists():
            universidad_seleccionada = universidades.first().id_universidad
        descuentos = DescuentoBanco.objects.filter(universidad_id=universidad_seleccionada).order_by('id')
    else:
        univ_coord = request.user.universidad_coordinador
        universidad_seleccionada = univ_coord.id_universidad if univ_coord else None
        descuentos = DescuentoBanco.objects.filter(universidad=univ_coord).order_by('id')
        if univ_coord:
            universidades = Universidad.objects.filter(id_universidad=univ_coord.id_universidad)

    base_template = 'base_admin.html' if request.user.is_staff else 'base_public.html'

    return render(request, 'admin_usuarios/admin_descuentos.html', {
        'descuentos': descuentos,
        'universidades': universidades,
        'universidad_seleccionada': universidad_seleccionada,
        'base_template': base_template
    })

@login_required
def admin_guardar_descuento(request):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')

    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        icono = request.POST.get('icono', 'fa-building-columns')
        univ_id_post = request.POST.get('universidad_id')

        if nombre:
            from user.models import DescuentoBanco
            from app.models import Universidad
            
            univ_asignar = request.user.universidad_coordinador
            if not univ_asignar and univ_id_post and univ_id_post.isdigit():
                univ_asignar = Universidad.objects.filter(id_universidad=univ_id_post).first()

            DescuentoBanco.objects.create(nombre=nombre, icono=icono, universidad=univ_asignar)
            messages.success(request, f'Banco "{nombre}" agregado correctamente.')

    url = reverse('admin_descuentos')
    if request.POST.get('universidad_id'): url += f"?universidad={request.POST.get('universidad_id')}"
    return redirect(url)

@login_required
def admin_eliminar_descuento(request, desc_id):
    if request.method == 'POST':
        from user.models import DescuentoBanco
        from django.shortcuts import get_object_or_404
        desc = get_object_or_404(DescuentoBanco, id=desc_id)
        desc.delete()
        messages.success(request, 'Banco eliminado.')
    url = reverse('admin_descuentos')
    if request.POST.get('universidad_id'): url += f"?universidad={request.POST.get('universidad_id')}"
    return redirect(url)


@login_required
def admin_membresias(request):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')

    from user.models import PlanMembresia, SuscripcionUsuario, Usuario
    from app.models import Universidad

    universidades = Universidad.objects.all()
    univ_id = request.GET.get('universidad')
    
    # Lógica Multi-Tenant
    if request.user.is_staff and not request.user.universidad_coordinador:
        universidad_seleccionada = Universidad.objects.filter(id_universidad=univ_id).first() if univ_id and univ_id.isdigit() else universidades.first()
    else:
        universidad_seleccionada = request.user.universidad_coordinador
        universidades = Universidad.objects.filter(id_universidad=universidad_seleccionada.id_universidad) if universidad_seleccionada else []

    # Obtenemos o creamos el plan para esta universidad
    plan, created = PlanMembresia.objects.get_or_create(universidad=universidad_seleccionada)

    # Actualizar el plan si se envió el formulario
    if request.method == 'POST' and request.POST.get('action') == 'update_plan':
        plan.valor_anual = request.POST.get('valor_anual', 0)
        plan.descripcion = request.POST.get('descripcion', '')
        plan.save()
        messages.success(request, "Plan de membresía actualizado correctamente.")
        return redirect(f"{reverse('admin_membresias')}?universidad={universidad_seleccionada.id_universidad}")

    # Obtenemos los usuarios de esta universidad para ver su estado
    # Asumiendo que puedes filtrar usuarios por universidad (ajusta este filtro a tu modelo real)
    usuarios = Usuario.objects.filter(carrera__departamento__facultad__universidad=universidad_seleccionada)
    
    base_template = 'base_admin.html' if request.user.is_staff else 'base_public.html'

    return render(request, 'admin_usuarios/admin_membresias.html', {
        'universidades': universidades,
        'universidad_seleccionada': universidad_seleccionada,
        'plan': plan,
        'usuarios': usuarios,
        'base_template': base_template
    })




@login_required
def admin_mensajeria(request):
    if not request.user.is_staff and not request.user.es_coordinador:
        messages.error(request, "No tienes permisos.")
        return redirect('home')

    from app.models import Universidad, Facultad, Carrera
    # Asegúrate de importar LogCorreo desde donde lo hayas creado (app.models o user.models)
    from app.models import LogCorreo 
    from user.models import Usuario

    # Multi-tenant: Qué puede ver este admin
    if request.user.is_staff and not request.user.universidad_coordinador:
        universidades = Universidad.objects.all()
    else:
        univ_coord = request.user.universidad_coordinador
        universidades = Universidad.objects.filter(id_universidad=univ_coord.id_universidad) if univ_coord else []

    facultades = Facultad.objects.all()
    carreras = Carrera.objects.all()

    # Historial de correos (Logs)
    logs = LogCorreo.objects.all().order_by('-fecha_envio')
    if not request.user.is_staff:
        logs = logs.filter(remitente=request.user)

    # --- PROCESAMIENTO DEL ENVÍO DE EMAIL MASIVO ---
    if request.method == 'POST':
        univ_id = request.POST.get('universidad')
        fac_id = request.POST.get('facultad')
        carr_id = request.POST.get('carrera')
        asunto = request.POST.get('asunto')
        cuerpo = request.POST.get('cuerpo')
        archivo = request.FILES.get('archivo')

        # 1. Filtramos usuarios activos que tengan email
        usuarios = Usuario.objects.filter(is_active=True).exclude(email="")
        filtros_texto = []

        if univ_id:
            usuarios = usuarios.filter(carrera__departamento__facultad__universidad_id=univ_id)
            from app.models import Universidad
            univ_nombre = Universidad.objects.filter(pk=univ_id).values_list('nombre', flat=True).first()
            filtros_texto.append(univ_nombre or f"Universidad {univ_id}")
        if fac_id:
            usuarios = usuarios.filter(carrera__departamento__facultad_id=fac_id)
            from app.models import Facultad
            fac_nombre = Facultad.objects.filter(pk=fac_id).values_list('nombre', flat=True).first()
            filtros_texto.append(fac_nombre or f"Facultad {fac_id}")
        if carr_id:
            usuarios = usuarios.filter(carrera_id=carr_id)
            from app.models import Carrera
            carr_nombre = Carrera.objects.filter(pk=carr_id).values_list('nombre', flat=True).first()
            filtros_texto.append(carr_nombre or f"Carrera {carr_id}")

        if not usuarios.exists():
            messages.warning(request, "No se encontraron usuarios con los filtros seleccionados.")
        else:
            try:
                mensajes_email = []
                
                # TRUCO PRO: Si hay archivo, lo leemos una sola vez en memoria. 
                # Si lo leemos dentro del 'for', el archivo se agota en el primer ciclo.
                adjunto_nombre = None
                adjunto_contenido = None
                adjunto_mime = None
                if archivo:
                    adjunto_nombre = archivo.name
                    adjunto_contenido = archivo.read()
                    adjunto_mime = archivo.content_type

                # 2. Armamos la lista de mensajes (Tu lógica)
                for usuario in usuarios:
                    email = EmailMessage(
                        subject=asunto,
                        # Opcional: si quieres, puedes personalizar el cuerpo aquí 
                        # usando body=cuerpo.replace("{nombre}", usuario.first_name)
                        body=cuerpo,
                        from_email=settings.DEFAULT_FROM_EMAIL,
                        to=[usuario.email]
                    )
                    
                    if archivo:
                        email.attach(adjunto_nombre, adjunto_contenido, adjunto_mime)
                        
                    mensajes_email.append(email)

                # 3. Enviamos todos de golpe abriendo una sola conexión
                if mensajes_email:
                    import smtplib
                    import ssl
                    from email.mime.multipart import MIMEMultipart
                    from email.mime.text import MIMEText
                    from email.mime.base import MIMEBase
                    from email import encoders

                    context = ssl.create_default_context()
                    context.check_hostname = False
                    context.verify_mode = ssl.CERT_NONE

                    enviados = 0
                    with smtplib.SMTP('smtp-relay.brevo.com', 587) as server:
                        server.ehlo()
                        server.starttls(context=context)
                        server.login('a4ba9b001@smtp-brevo.com', 'UPLM92nEtK50FTcD')
                        
                        for usuario in usuarios:
                            try:
                                msg = MIMEMultipart()
                                msg['Subject'] = asunto
                                msg['From'] = settings.DEFAULT_FROM_EMAIL
                                msg['To'] = usuario.email
                                msg.attach(MIMEText(cuerpo, 'plain'))
                                
                                if adjunto_contenido:
                                    part = MIMEBase('application', 'octet-stream')
                                    part.set_payload(adjunto_contenido)
                                    encoders.encode_base64(part)
                                    part.add_header('Content-Disposition', f'attachment; filename="{adjunto_nombre}"')
                                    msg.attach(part)
                                
                                server.sendmail(settings.DEFAULT_FROM_EMAIL, usuario.email, msg.as_string())
                                enviados += 1
                            except Exception as e:
                                print(f"Error enviando a {usuario.email}: {e}")
                    
                    # 4. Guardamos el Log en la Base de Datos
                    LogCorreo.objects.create(
                        remitente=request.user,
                        asunto=asunto,
                        cuerpo=cuerpo,
                        archivo_adjunto=archivo, # Se guarda el archivo en el servidor
                        filtros_aplicados=" > ".join(filtros_texto) if filtros_texto else "Todos los usuarios",
                        cantidad_destinatarios=enviados # Guardamos el número real devuelto por la conexión
                    )
                    
                    messages.success(request, f"¡Éxito! Se enviaron {enviados} correos masivos de forma segura.")
            
            except Exception as e:
                print(f"Error Anymail (Mensajería Masiva): {e}")
                messages.error(request, f"Error al enviar correos: {str(e)}")
        
        return redirect('admin_mensajeria')

    base_template = 'base_admin.html' if request.user.is_staff else 'base_public.html'

    return render(request, 'admin_usuarios/admin_mensajeria.html', {
        'universidades': universidades,
        'facultades': facultades,
        'carreras': carreras,
        'logs': logs,
        'base_template': base_template
    })

@login_required
@user_passes_test(lambda u: u.is_staff)
def instituciones_coordinadores(request):
    universidades = Universidad.objects.all()
    
    data = []
    for uni in universidades:
        coordinador = Usuario.objects.filter(
            universidad_coordinador=uni,
            es_coordinador=True
        ).first()
        data.append({
            'universidad': uni,
            'coordinador': coordinador
        })
    
    return render(request, 'admin_usuarios/instituciones_coordinadores.html', {
        'data': data
    })


@login_required
@user_passes_test(lambda u: u.is_staff)
def eliminar_log_correo(request, log_id):
    from app.models import LogCorreo
    log = get_object_or_404(LogCorreo, pk=log_id)
    log.delete()
    messages.success(request, "Registro eliminado.")
    return redirect('admin_mensajeria')

@login_required
@user_passes_test(lambda u: u.is_staff)
def eliminar_todos_logs(request):
    from app.models import LogCorreo
    LogCorreo.objects.all().delete()
    messages.success(request, "Historial eliminado completamente.")
    return redirect('admin_mensajeria')


# TIENDA
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from app.models import Producto, OrdenCompra, Universidad


@login_required
def gestion_tienda(request):
    if not request.user.is_staff and (not request.user.es_coordinador or not request.user.universidad_coordinador):
        return redirect('home')

    if request.user.is_staff:
        productos = Producto.objects.all().order_by('-creado')
        ordenes = OrdenCompra.objects.all().order_by('-creado')[:10]
        universidad = None
    else:
        productos = Producto.objects.filter(
            universidad=request.user.universidad_coordinador
        ).order_by('-creado')

        ordenes = OrdenCompra.objects.filter(
            producto__universidad=request.user.universidad_coordinador
        ).order_by('-creado')[:10]

        universidad = request.user.universidad_coordinador

    return render(request, 'admin_usuarios/gestion_tienda.html', {
        'productos': productos,
        'ordenes': ordenes,
        'universidad': universidad
    })


@login_required
def crear_producto(request):
    if not request.user.is_staff and (not request.user.es_coordinador or not request.user.universidad_coordinador):
        return redirect('home')

    if request.method == 'POST':
        print("FILES:", request.FILES)
        print("POST:", request.POST)
        nombre = request.POST.get('nombre')
        descripcion = request.POST.get('descripcion')
        precio = request.POST.get('precio')
        stock = request.POST.get('stock')
        imagen = None
        if request.FILES.get('imagen'):
            upload_result = cloudinary.uploader.upload(
                request.FILES.get('imagen'),
                transformation=[{"width": 800, "crop": "limit", "quality": "auto"}]
            )
            imagen = upload_result.get('secure_url')

        if nombre and precio and stock:
            if request.user.is_staff:
                uni_id = request.POST.get('universidad_id')
                universidad_obj = Universidad.objects.filter(pk=uni_id).first()
            else:
                universidad_obj = request.user.universidad_coordinador

            Producto.objects.create(
                universidad=universidad_obj,
                creado_por=request.user,
                nombre=nombre,
                descripcion=descripcion,
                precio=int(precio),
                stock=int(stock),
                imagen=imagen,
                activo=True
            )

            messages.success(request, "Producto creado exitosamente.")
        else:
            messages.error(request, "Completa todos los campos obligatorios.")

        return redirect('gestion_tienda')

    universidades = Universidad.objects.all()
    return render(request, 'admin_usuarios/crear_producto.html', {
        'universidades': universidades
    })


@login_required
def eliminar_producto(request, producto_id):
    producto = get_object_or_404(Producto, pk=producto_id)

    if request.user == producto.creado_por or request.user.is_staff:
        producto.delete()
        messages.success(request, "Producto eliminado.")

    return redirect('gestion_tienda')


@login_required
def editar_producto(request, producto_id):
    producto = get_object_or_404(Producto, pk=producto_id)

    if request.method == 'POST':
        producto.nombre = request.POST.get('nombre', producto.nombre)
        producto.descripcion = request.POST.get('descripcion', producto.descripcion)
        producto.precio = int(request.POST.get('precio', producto.precio))
        producto.stock = int(request.POST.get('stock', producto.stock))
        producto.activo = request.POST.get('activo') == 'on'

        if request.FILES.get('imagen'):
            producto.imagen = None
        if request.FILES.get('imagen'):
            upload_result = cloudinary.uploader.upload(request.FILES.get('imagen'))
            imagen = upload_result.get('secure_url')

        producto.save()
        messages.success(request, "Producto actualizado.")
        return redirect('gestion_tienda')

    return render(request, 'admin_usuarios/editar_producto.html', {
        'producto': producto
    })
    
# ============================================
# ENCUESTAS
# ============================================
from app.models import Encuesta, PreguntaEncuesta, OpcionPregunta, RespuestaEncuesta, DetalleRespuesta

@login_required
def gestion_encuestas(request):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')

    if request.user.is_staff:
        encuestas = Encuesta.objects.all().order_by('-creada')
    else:
        encuestas = Encuesta.objects.filter(
            universidad=request.user.universidad_coordinador
        ).order_by('-creada')

    base_template = 'base_admin.html' if request.user.is_staff else 'base_public.html'
    return render(request, 'encuestas/gestion_encuestas.html', {
        'encuestas': encuestas,
        'base_template': base_template
    })

@login_required
def crear_encuesta(request):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')
    
    # AJAX para filtrar carreras por universidad
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        action = request.POST.get('action')
        if action == 'buscar_carreras_por_universidad':
            uni_id = request.POST.get('id')
            carreras = Carrera.objects.filter(
                departamento__facultad__universidad_id=uni_id
            ).values('pk', 'nombre').order_by('nombre')
            return JsonResponse(list(carreras), safe=False)
        
    if request.user.is_staff:
        universidades = Universidad.objects.all()
    else:
        universidades = [request.user.universidad_coordinador]

    if request.method == 'POST':
        titulo = request.POST.get('titulo')
        universidad_id = request.POST.get('universidad_id') if request.user.is_staff else request.user.universidad_coordinador.pk
        carrera_id = request.POST.get('carrera_id') or None
        fecha_vencimiento = request.POST.get('fecha_vencimiento')

        universidad_obj = Universidad.objects.get(pk=universidad_id)
        carrera_obj = Carrera.objects.filter(pk=carrera_id).first() if carrera_id else None

        encuesta = Encuesta.objects.create(
            titulo=titulo,
            universidad=universidad_obj,
            carrera=carrera_obj,
            creada_por=request.user,
            fecha_vencimiento=fecha_vencimiento,
            activa=True
        )

        # Guardar preguntas de opción múltiple
        i = 0
        while True:
            texto_pregunta = request.POST.get(f'pregunta_{i}')
            if texto_pregunta is None:
                break
            pregunta = PreguntaEncuesta.objects.create(
                encuesta=encuesta,
                texto=texto_pregunta,
                tipo='opcion_multiple',
                orden=i
            )
            j = 0
            while True:
                opcion = request.POST.get(f'opcion_{i}_{j}')
                if opcion is None:
                    break
                OpcionPregunta.objects.create(pregunta=pregunta, texto=opcion)
                j += 1
            i += 1

        # Guardar pregunta de texto
        texto_libre = request.POST.get('pregunta_texto')
        if texto_libre:
            PreguntaEncuesta.objects.create(
                encuesta=encuesta,
                texto=texto_libre,
                tipo='texto',
                orden=i
            )

        messages.success(request, "Encuesta creada exitosamente.")
        return redirect('gestion_encuestas')

    base_template = 'base_admin.html' if request.user.is_staff else 'base_public.html'

    if request.user.is_staff:
        carreras = Carrera.objects.all()
    else:
        carreras = Carrera.objects.filter(
            departamento__facultad__universidad=request.user.universidad_coordinador
        )

    return render(request, 'encuestas/crear_encuesta.html', {
        'universidades': universidades,
        'carreras': carreras,
        'base_template': base_template
    })


@login_required
def ver_resultados_encuesta(request, encuesta_id):
    if not request.user.is_staff and not request.user.es_coordinador:
        return redirect('home')

    encuesta = get_object_or_404(Encuesta, pk=encuesta_id)
    preguntas = encuesta.preguntas.all().order_by('orden')
    total_respuestas = RespuestaEncuesta.objects.filter(encuesta=encuesta).count()

    resultados = []
    for pregunta in preguntas:
        if pregunta.tipo == 'opcion_multiple':
            opciones = []
            for opcion in pregunta.opciones.all():
                count = DetalleRespuesta.objects.filter(
                    pregunta=pregunta,
                    opcion_seleccionada=opcion
                ).count()
                porcentaje = round((count / total_respuestas * 100), 1) if total_respuestas > 0 else 0
                opciones.append({
                    'texto': opcion.texto,
                    'count': count,
                    'porcentaje': porcentaje
                })
            resultados.append({
                'pregunta': pregunta,
                'tipo': 'opcion_multiple',
                'opciones': opciones
            })
        else:
            respuestas_texto = DetalleRespuesta.objects.filter(
                pregunta=pregunta
            ).exclude(texto_respuesta='').values_list('texto_respuesta', flat=True)
            resultados.append({
                'pregunta': pregunta,
                'tipo': 'texto',
                'respuestas': list(respuestas_texto)
            })

    base_template = 'base_admin.html' if request.user.is_staff else 'base_public.html'
    return render(request, 'encuestas/resultados_encuesta.html', {
        'encuesta': encuesta,
        'resultados': resultados,
        'total_respuestas': total_respuestas,
        'base_template': base_template
    })


@login_required
def descargar_resultados_encuesta(request, encuesta_id):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from django.http import HttpResponse

    encuesta = get_object_or_404(Encuesta, pk=encuesta_id)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Resultados"

    # Estilos
    header_font = Font(name='Arial', bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill('solid', start_color='1e3a5f')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    data_font = Font(name='Arial', size=10)
    data_align = Alignment(vertical='center', wrap_text=True)

    # Fila de título
    preguntas = list(encuesta.preguntas.all().order_by('orden'))
    total_cols = 2 + len(preguntas)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=total_cols)
    titulo_cell = ws.cell(row=1, column=1, value=f"Resultados: {encuesta.titulo}")
    titulo_cell.font = Font(name='Arial', bold=True, color='FFFFFF', size=13)
    titulo_cell.fill = PatternFill('solid', start_color='991b1b')
    titulo_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30

    # Fila de info
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=total_cols)
    info_cell = ws.cell(row=2, column=1, value=f"Sede: {encuesta.universidad.nombre} | Vence: {encuesta.fecha_vencimiento} | Total respuestas: {RespuestaEncuesta.objects.filter(encuesta=encuesta).count()}")
    info_cell.font = Font(name='Arial', italic=True, size=9, color='555555')
    info_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[2].height = 18

    # Encabezados
    headers = ['Usuario', 'Email'] + [p.texto for p in preguntas]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = border
    ws.row_dimensions[3].height = 40

    # Datos
    for row_idx, respuesta in enumerate(RespuestaEncuesta.objects.filter(encuesta=encuesta).select_related('usuario'), 4):
        fill = PatternFill('solid', start_color='EFF3FB') if row_idx % 2 == 0 else PatternFill('solid', start_color='FFFFFF')
        fila = [respuesta.usuario.username, respuesta.usuario.email]
        for pregunta in preguntas:
            detalle = DetalleRespuesta.objects.filter(respuesta=respuesta, pregunta=pregunta).first()
            if detalle:
                if pregunta.tipo == 'opcion_multiple' and detalle.opcion_seleccionada:
                    fila.append(detalle.opcion_seleccionada.texto)
                else:
                    fila.append(detalle.texto_respuesta or '')
            else:
                fila.append('')
        for col, valor in enumerate(fila, 1):
            cell = ws.cell(row=row_idx, column=col, value=valor)
            cell.font = data_font
            cell.alignment = data_align
            cell.fill = fill
            cell.border = border
        ws.row_dimensions[row_idx].height = 20

    # Anchos de columna
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 30
    for col in range(3, total_cols + 1):
        ws.column_dimensions[get_column_letter(col)].width = 35

    # Freeze headers
    ws.freeze_panes = 'A4'

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="resultados_{encuesta.titulo}.xlsx"'
    wb.save(response)
    return response


@login_required
def eliminar_encuesta(request, encuesta_id):
    encuesta = get_object_or_404(Encuesta, pk=encuesta_id)
    if request.user.is_staff or encuesta.creada_por == request.user:
        encuesta.delete()
        messages.success(request, "Encuesta eliminada.")
    return redirect('gestion_encuestas')